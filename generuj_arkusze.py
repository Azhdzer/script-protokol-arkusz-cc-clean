# -*- coding: utf-8 -*-
"""
generuj_arkusze.py

Automatyczne tworzenie kopii arkuszy obliczeniowych na podstawie pliku protokołu.

Etap 1 – tworzenie nazwanych kopii szablonu
Etap 2 – zarządzanie zakładkami (zmiana nazw, usuwanie/kopiowanie)
Etap 3 – wypełnianie C15:C19 i D15:D19
         (dla CC: dane stałe z kol. L/M; dla CC-04: dynamiczne kolumny C/D)
Etap 4 – wypełnianie E15:E19 i F15:F19 (dane zmienne per kopia, kol. Q/R i dalej)
Etap 5 – wypełnianie komórek nagłówkowych i stopkowych
         (E4:F4, G6, E5, E6, H57, B228, H228, B230:C230, H230:I230)
Etap 6 – wypełnianie arkusza Wyniki: F24 (per kopia z Strona 3 Q17/S17/…),
         C28, C32 (daty), E28:G28, E32:G32 (podpisy)
Etap 7 – tworzenie kopii dokumentów Word (świadectwa wzorcowania):
         podmiana placeholderów, tabela kalibracyjna D246–G246 per zakładka

UWAGA: openpyxl odczytuje wartości zapisane przez Excel (data_only=True).
Jeśli plik protokołu nie był przeliczony i zapisany w Excelu, niektóre komórki
mogą zwrócić None zamiast wartości formuły.
"""

import os
import re
import time
import shutil
import zipfile
import datetime
import math
import urllib.parse
from copy import deepcopy
from itertools import groupby
import openpyxl
import xlwings as xw

import pz_dane   # wspolny modul: dane przyrzadow z PZ (fallback, gdy Strona 2 pusta)
import cc_config as C   # rejestr ustawien + odczyt zmiennych srodowiskowych z panelu

# ---------------------------------------------------------------------------
# Wizualne helpers logowania
# ---------------------------------------------------------------------------
def _warn(msg, indent="  "):
    """Ostrzezenie — wyroznia sie w logu, ale wykonanie trwa dalej."""
    print(f"{indent}{'!':->3} UWAGA {'!':->3}  {msg}")

def _err(msg, indent="  "):
    """Blad krytyczny — akcja nie wykonala sie."""
    sep = "!" * max(60, len(msg) + 12)
    print(f"{indent}{sep}")
    print(f"{indent}!!! BLAD: {msg}")
    print(f"{indent}{sep}")

def _ok(msg, indent="  "):
    """Potwierdzenie pomyslnej operacji."""
    print(f"{indent}[OK] {msg}")

def _log_etap(msg, t0, indent="      "):
    """
    Log postepu z czasem od poczatku operacji na kopii. flush=True jest KLUCZOWE —
    dzieki temu widac ostatni wykonany krok nawet gdy Excel sie zawiesi.
    """
    print(f"{indent}[{time.time() - t0:6.1f}s] {msg}", flush=True)


def _zamknij_sesje_excel(app, linked_wbs):
    """Zamyka pliki linkowane i konczy proces Excela (odporne na pad Excela)."""
    for lwb in (linked_wbs or []):
        try:
            lwb.close()
        except Exception:
            pass
    if app is not None:
        try:
            app.quit()
        except Exception as _eq:
            _warn(f"app.quit() nie powiodlo sie (Excel prawdopodobnie juz padl): "
                  f"{type(_eq).__name__}: {_eq}")


def _nowa_sesja_excel(folder):
    """
    Tworzy SWIEZA instancje Excela i otwiera w niej pliki linkowane (Wzory / Obliczenia).
    Zwraca (app, linked_wbs, sciezki_linkowane).

    Kazda kopia dostaje wlasna sesje: kopiowanie zakladek z obiektami OLE
    (Equation/Word.Document) zostawia w procesie Excela zasoby, ktorych on nie zwalnia.
    Przy kilku kopiach z wieloma punktami konczylo sie to padem Excela (RPC -2147023170).
    Restart procesu miedzy kopiami kosztuje kilka sekund i eliminuje kumulacje.
    """
    app = xw.App(visible=False, add_book=False)
    app.api.AutomationSecurity = 1   # msoAutomationSecurityLow — wlacza makra bez pytania
    app.api.DisplayAlerts = False
    app.api.AskToUpdateLinks = False
    app.api.Visible = False
    # Blokuj Workbook_Open/Worksheet_Calculate makr JUZ PRZED otwarciem plikow linkowanych
    # (ich makra siegaja \\plum4 i moga uruchamiac Worda).
    app.api.EnableEvents = False
    try:
        app.api.AutoRecover.Enabled = False  # nie twórz plików autoodzyskiwania (.xar)
    except Exception:
        pass

    linked_wbs = []
    sciezki_linkowane = {}
    for plik_link in PLIKI_LINKOWANE:
        sciezka_link = os.path.join(folder, plik_link)
        if os.path.exists(sciezka_link):
            try:
                lwb = _open_book_hidden(app, sciezka_link, update_links=False)
                linked_wbs.append(lwb)
                sciezki_linkowane[plik_link.lower()] = lwb.fullname
            except Exception as exc:
                _warn(f"Nie mozna otworzyc pliku linkowanego: {plik_link}\n"
                      f"      Sciezka lokalna: {sciezka_link}\n"
                      f"      Blad: {type(exc).__name__}: {exc}")
        else:
            _warn(f"Brak pliku linkowanego w folderze roboczym: {plik_link}\n"
                  f"      Sciezka serwerowa: {_info_serwer(plik_link)}\n"
                  f"      Oczekiwano w: {sciezka_link}")
    return app, linked_wbs, sciezki_linkowane

def _open_book_hidden(app, path, **kwargs):
    """
    Otwiera skoroszyt i przywraca ustawienia COM po evencie Workbook_Open.
    Makra moga resetowac Visible/DisplayAlerts — ustawiamy je z powrotem.
    """
    wb = app.books.open(path, **kwargs)
    try:
        app.api.Visible = False
        app.api.DisplayAlerts = False
        app.api.AskToUpdateLinks = False
    except Exception:
        pass
    return wb

def _info_serwer(nazwa_pliku):
    """Zwraca sciezke serwerowa dla danego pliku (do wyswietlenia w logach)."""
    return LINKI_SERWEROWE.get(nazwa_pliku, LINKI_SERWEROWE.get(nazwa_pliku.lower(), "nieznana"))
# ---------------------------------------------------------------------------

try:
    from docx import Document as DocxDocument
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    _DOCX_OK = True
except ImportError:
    DocxDocument = None  # type: ignore
    _DOCX_OK = False

# =============================================================================
# KONFIGURACJA
#
# Wszystkie wartosci ponizej ustawia sie w PANELU (app_gui.py) — to, co widzisz
# w kodzie, to tylko wartosci DOMYSLNE uzywane przy recznym uruchomieniu skryptu.
# Panel podaje je przez zmienne srodowiskowe (nazwy w nawiasach przy C.*).
# =============================================================================

FOLDER           = os.environ.get("CC_FOLDER") or \
                   r"C:\Users\artisom.azhdzer\Desktop\Script protokoł - arkusz CC"   # folder z plikami xlsx; "." = ten sam co skrypt  r"."
                           # możesz podać pełną ścieżkę, np. r"C:\Moje\Pliki"

PROTOKOL_PLIK    = os.environ.get("CC_PROTOKOL") or "188_LA_TH_2026 - protokół CC_2.xlsx"
SZABLON_PLIK     = os.environ.get("CC_SZABLON") or "xxx_LA_TH_2026 - ILAJ 5.4_11#21 - Wzór ark. obl.Wer.12 z 17.06.2026 - 1 - RH (CC).xlsx"

# Uklad protokolu i szablonu opisuja stale w sekcji "UKLAD FORMULARZY" ponizej
# (nazwy arkuszy, wiersze, kolumny) — to nie sa opcje do zmiany.
PODPISUJACY_1    = C.tekst("GEN_PODPIS_1", "Artsiom Azhdzer")   # B230:C230 (scalona) — podpisujacy z lewej
PODPISUJACY_2    = C.tekst("GEN_PODPIS_2", "Marek Szpakowski")  # H230:I230 (scalona) — podpisujacy z prawej

# --- K18: higrometr punktu rosy w kazdej kopii-zakladce -----------------------
# Wartosc zalezy od KOMORY, w ktorej wykonano wzorcowanie, dlatego ustawia sie ja
# osobno dla kazdej komory. Klucz to typ protokolu:
#   "CC"    — komora CC        (protokol '... - protokół CC.xlsx')
#   "CC-04" — komora CC-04     (protokol '... - protokół CC-04.xlsx')
# Typowe wartosci: "S8000-02", "S8000", "OPTIDEW", "-".
#
# UWAGA: dla punktow TYLKO-TEMPERATURA (zakladka bez RH, np. nazwa '5, -') w K18
# ZAWSZE wpisywane jest "-", niezaleznie od tego ustawienia.
# Arkusza "Wyniki" to nie dotyczy (nie jest modyfikowany).
HIGROMETR_K18_WG_KOMORY = {
    "CC":    C.tekst("GEN_K18_CC", "S8000-02"),
    "CC-04": C.tekst("GEN_K18_CC04", "S8000"),
}
# Uzywane, gdy typ komory nie figuruje w slowniku powyzej.
HIGROMETR_K18_DOMYSLNY = C.tekst("GEN_K18_DOM", "S8000")


def _higrometr_k18(is_cc04):
    """Higrometr do K18 dla biezacego protokolu (CC albo CC-04)."""
    return HIGROMETR_K18_WG_KOMORY.get("CC-04" if is_cc04 else "CC",
                                       HIGROMETR_K18_DOMYSLNY)

# szablon Word, gdy brak aktywnej wilgotnosci
SZABLON_WORD_TYLKO_TEMP = C.tekst("GEN_WORD_TEMP", "xxx_yyy_LA_TH_2026 - tylko temp.docx")
# szablon Word, gdy WSZYSTKIE zakladki maja aktywna wilgotnosc
SZABLON_WORD_Z_RH       = C.tekst("GEN_WORD_RH", "xxx_yyy_LA_TH_2026 - zakres.docx")
# szablon Word, gdy CZESC zakladek ma wilgotnosc, a czesc nie (dwie tabele)
SZABLON_WORD_MIESZANY   = C.tekst("GEN_WORD_MIX", "xxx_yyy_LA_TH_2026 - zakres + temp.docx")
NR_SW_POCZATKOWY    = C.calk("GEN_NR_SW", 1047)   # numer świadectwa pierwszej kopii (rosnaco dla kolejnych)

NR_POMIESZCZENIA = C.calk("GEN_NR_POM", 9)                  # numer pomieszczenia środowiskowego
MODEL_CZUJNIKA   = C.tekst("GEN_MODEL_CZUJ", "MX1101-02")   # model czujnika środowiskowego

# Sterowanie etapami:
# - GENERUJ_EXCEL=False: nie tworzy/nie modyfikuje kopii Excel,
#                       korzysta z juz istniejacych kopii (gdy GENERUJ_WORD=True).
# - GENERUJ_WORD=False : pomija Etap 7 (Word).
GENERUJ_EXCEL = True
GENERUJ_WORD  = True

# Pliki linkowane wymagane do przeliczenia formul kalibracyjnych (D246/F246/G246).
# Muszą być otwarte w tej samej sesji Excel — podaj dokładne nazwy z rozszerzeniem.
PLIKI_LINKOWANE     = C.lista("GEN_LINKOWANE", [
    "Obliczenia tdp, RH, C.xls",
    "Wzory.xls",
])

# Od ilu kopiowanych zakladek ostrzegac w logu, ze potrwa to dlugo. Czas jednej kopii
# NIE jest zakladany z gory — mierzymy pierwsza i na tej podstawie szacujemy reszte
# (zalezy od maszyny i od liczby obiektow OLE w szablonie).
PROG_OSTRZEZENIA_KOPII = C.calk("GEN_PROG_OSTRZ", 10)

# Szerokosc paska zakladek w zapisanej kopii (0.0-1.0; reszta miejsca idzie na poziomy
# pasek przewijania). Domyslne 0.6 Excela bywa za male przy wielu punktach i zakladki
# chowaja sie za strzalkami — 0.85 sprawia, ze po otwarciu widac cala liste.
TAB_RATIO = C.liczba("GEN_TAB_RATIO", 0.85)

# Warunki srodowiskowe (Pom. nr 9): maksymalna odleglosc czasowa rekordu czujnika od
# punktu pomiarowego. Gdy najblizszy rekord jest dalej — F/G zostaja PUSTE zamiast
# wpisywac warunki z zupelnie innej chwili (wczesniej brany byl ostatni rekord z pliku,
# przez co kilka punktow dostawalo te same wartosci).
# Czujnik Pom. nr 9 zapisuje co 60 s bez przerw (sprawdzone na pliku rocznym),
# wiec najblizszy rekord jest zawsze w granicach ~30 s. Zapas 2 min pokrywa
# ewentualna zmiane interwalu, a jednoczesnie nie pozwala wpisac warunkow
# sprzed pol godziny. Gdy nic nie miesci sie w tolerancji, F/G zostaja PUSTE.
TOLERANCJA_CZUJNIKA_MIN = C.liczba("GEN_TOL_CZUJ", 2.0)

# =============================================================================
# UKLAD FORMULARZY — stale wynikajace z wzorow PLUM (ILAJ 5.4/11).
# To NIE sa opcje do zmiany: inna wartosc wymagalaby przerobienia samych formularzy.
# Trzymane osobno, zeby nie mieszaly sie z konfiguracja powyzej.
# =============================================================================
ARKUSZ_STRONA2 = "Strona 2"   # protokol: tabela przyrzadow (= lista kopii do zrobienia)
ARKUSZ_STRONA3 = "Strona 3"   # protokol: punkty pomiarowe (= zakladki w kopii)
# Arkusz zbiorczy w kopii. NIE jest zakladka punktu pomiarowego, wiec nie podlega
# usuwaniu/przemianowaniu razem z nimi — ale skrypt go modyfikuje (tabele wynikow,
# formula histerezy, podpisy, F24/C28/C32).
ARKUSZ_WYNIKI  = "Wyniki"

START_ROW_S2   = 11   # pierwszy wiersz przyrzadu w Strona 2
START_ROW_S3   = 20   # pierwszy wiersz bloku punktu w Strona 3
BLOK_S3        = 5    # wierszy na jeden punkt pomiarowy w Strona 3

START_COL_E_S3 = 17   # kolumna Q — odczyty 1. przyrzadu (zrodlo dla E15:E19 w kopii)
START_COL_F_S3 = 18   # kolumna R — wilgotnosc 1. przyrzadu (zrodlo dla F15:F19)
KROK_COL_EF    = 2    # kazdy kolejny przyrzad to para kolumn dalej w prawo

# Docelowe sciezki serwerowe dla linkow zewnetrznych, ktore maja byc
# przywrocone na koncu (po wypelnieniu i odczycie kalibracji).
LINKI_SERWEROWE = {
    "Obliczenia tdp, RH, C.xls": C.tekst(
        "GEN_LINK_OBLICZENIA", r"\\plum4\LabPomiarowe\Obliczenia tdp, RH, C.xls"),
    "Wzory.xls": C.tekst("GEN_LINK_WZORY", r"\\plum4\LabPomiarowe\Wzory.xls"),
}

# --- Sprzatanie plikow autoodzyskiwania Excela (%AppData%\Microsoft\Excel) ---
# Te pliki (.xar + migawki .xlsb/.xls) zostaja po awariach Excela i potrafia sie
# nazbierac do setek MB, przez co Excel przy starcie probuje je odzyskac i pada.
# Skrypt czysci je PRZED uruchomieniem, ale tylko starsze niz prog ponizej —
# dzieki temu ewentualne dzisiejsze odzyskiwanie RECZNEJ pracy nie zostanie usuniete.
CZYSC_AUTORECOVER = True              # False = nie ruszaj folderu autoodzyskiwania (patrz GEN_AUTOREC nizej)
CZYSC_AUTORECOVER_STARSZE_NIZ_DNI = C.calk("GEN_AUTOREC_DNI", 1)  # usuwaj tylko starsze niz tyle dni; 0 = czysc wszystko

# Dla protokolow CC-04 dane E/F startuja od S/T zamiast Q/R.
PRZESUNIECIE_STARTU_KOL_CC04 = 2
WIERSZ_TYPU_CC04_S3 = 14

# Mapowanie typu z S14:T14 (kolejne kopie: U14:V14, W14:X14, ...)
# na stale wartosci zapisywane do K11/K12/K13/K17 w zakladkach roboczych.
_MAP_CC04_DOMYSLNA = [
    ["LG", "Pt100-09", "1586A-02", "101", "CC-04-LG"],
    ["LD", "Pt100-01", "1586A-02", "105", "CC-04-LD"],
    ["PD", "Pt100-18", "1586A-02", "107", "CC-04-PD"],
    ["PG", "Pt100-13", "1586A-02", "103", "CC-04-PG"],
]


def _mapowanie_cc04(wiersze):
    """Zamienia tabele [tag, K11, K12, K13, K17] na slownik uzywany w Etapie 5."""
    mapa = {}
    for w in wiersze:
        if not w or not str(w[0]).strip():
            continue
        pola = [str(x).strip() if x is not None else "" for x in w] + [""] * 5
        mapa[pola[0].upper()] = {"K11": pola[1], "K12": pola[2],
                                 "K13": pola[3], "K17": pola[4]}
    return mapa


MAPOWANIE_TYPU_CC04 = _mapowanie_cc04(C.tabela("GEN_MAP_CC04", _MAP_CC04_DOMYSLNA))

# Dla CC-04:
# - C15:C19 bierzemy z kolumny zależnej od typu (LG/PG/LD/PD),
# - D15:D19 bierzemy z kolumny O (15).
MAPOWANIE_KOLUMNY_C_CC04 = {
    "LG": 11,  # K
    "PG": 12,  # L
    "LD": 13,  # M
    "PD": 14,  # N
}
KOLUMNA_D_CC04_S3 = 15  # O

# Filtr kolorow danych z arkusza Strona 3:
# - komorki zielone (#CCFFCC) sa brane,
# - komorki szare (#BFBFBF) sa pomijane,
# - pozostale kolory wg BIERZ_INNE_KOLORY_S3.
FILTRUJ_KOLOR_S3      = C.flaga("GEN_FILTR_KOLOR", True)
KOLOR_AKTYWNY_S3      = C.tekst("GEN_KOLOR_AKT", "#CCFFCC")
KOLOR_POMIJANY_S3     = C.tekst("GEN_KOLOR_POM", "#BFBFBF")
BIERZ_INNE_KOLORY_S3  = C.flaga("GEN_INNE_KOLORY", False)
# Jesli True, kopia dostaje tylko te zakladki, ktore maja dane E/F
# (po filtracji kolorow) dla konkretnej kopii.
USUWAJ_PUSTE_BLOKI_KOPII_S3 = C.flaga("GEN_PUSTE", True)

# Przyrzad wyszarzony w CALOSCI na Stronie 3 (zaden blok E/F nie jest aktywny)
# nie dostaje ani kopii Excel, ani swiadectwa Word. Bez tego powstawal plik z
# samym arkuszem Wyniki i swiadectwo z pusta tabela kalibracji.
POMIJAJ_PRZYRZADY_BEZ_DANYCH = C.flaga("GEN_POMIJAJ_PUSTE", True)


# --- Nadpisania z panelu GUI (app_gui.py) przez zmienne srodowiskowe ---
GENERUJ_EXCEL = C.flaga("GEN_EXCEL", GENERUJ_EXCEL)
GENERUJ_WORD  = C.flaga("GEN_WORD", GENERUJ_WORD)
CZYSC_AUTORECOVER = C.flaga("GEN_AUTOREC", CZYSC_AUTORECOVER)

# =============================================================================
# WARUNKI ŚRODOWISKOWE — czujnik + Wzory.xls
# =============================================================================

def _sciezka_pliku_czujnika(nr_pom, model, rok):
    """Returns path to sensor file: local folder first, then network."""
    plik = f"Pom. nr {nr_pom} ({model}) - {rok}.xlsx"
    lokalna = os.path.join(FOLDER, plik)
    if os.path.exists(lokalna):
        return lokalna
    folder_siec = (
        rf"\\PLUM4\LabPomiarowe\Odczyty z rejestratorów środowiskowych"
        rf"\Pom. nr {nr_pom} (dom. {model})"
    )
    return os.path.join(folder_siec, plik)


def _wczytaj_dane_czujnika_miesiac(sciezka_pliku, rok, miesiac):
    """Reads (datetime, temp, hum) rows from sensor xlsx for given month/year sheet."""
    if not os.path.exists(sciezka_pliku):
        print(f"    [UWAGA] Brak pliku czujnika srodowiskowego: {sciezka_pliku}")
        return []
    try:
        wb_c = openpyxl.load_workbook(sciezka_pliku, read_only=True, data_only=True)
        nazwa_ark = f"{miesiac:02d}'{str(rok)[2:]}"
        if nazwa_ark not in wb_c.sheetnames:
            print(f"    [UWAGA] Brak arkusza '{nazwa_ark}' w pliku czujnika. Dostepne: {wb_c.sheetnames}")
            wb_c.close()
            return []
        ws_c = wb_c[nazwa_ark]
        wyniki = []
        for row in ws_c.iter_rows(min_row=2, values_only=True):
            if not row or len(row) < 3 or row[0] is None:
                continue
            a_val, b_val, c_val = row[0], row[1], row[2]
            if isinstance(a_val, datetime.datetime):
                dt = a_val
            elif isinstance(a_val, str):
                try:
                    dt = datetime.datetime.strptime(a_val.strip(), "%d.%m.%Y %H:%M:%S")
                except ValueError:
                    continue
            else:
                continue
            try:
                temp = float(b_val) if b_val is not None else None
                hum = float(c_val) if c_val is not None else None
            except (TypeError, ValueError):
                continue
            if temp is not None:
                wyniki.append((dt, temp, hum))
        wb_c.close()
        return wyniki
    except Exception as e:
        print(f"    [UWAGA] Blad odczytu pliku czujnika: {e}")
        return []


def _szukaj_th_w_danych_czujnika(dane, target_dt, tol_min=None):
    """
    Zwraca (temp, hum) z rekordu czujnika NAJBLIZSZEGO czasowo target_dt — ale tylko gdy
    miesci sie w tolerancji TOLERANCJA_CZUJNIKA_MIN.

    Bez tego ograniczenia, gdy czujnik nie ma danych na dany czas (np. wzorcowanie trwalo
    dluzej niz zapis czujnika), zwracany byl OSTATNI rekord z pliku — i te same warunki
    srodowiskowe trafialy do wielu roznych punktow. Teraz w takim wypadku zwracamy
    (None, None), a komorki F/G zostaja puste do recznego uzupelnienia.
    Zwraca dodatkowo odchylke w sekundach jako trzeci element (do logu).
    """
    if not dane:
        return None, None, None
    if tol_min is None:
        tol_min = TOLERANCJA_CZUJNIKA_MIN
    best = min(dane, key=lambda x: abs((x[0] - target_dt).total_seconds()))
    odchylka_s = abs((best[0] - target_dt).total_seconds())
    if odchylka_s > tol_min * 60:
        return None, None, odchylka_s
    return best[1], best[2], odchylka_s


def _parse_date_s3(val):
    """Parses date from Strona 3 E column value (datetime.datetime, date, or string)."""
    if isinstance(val, datetime.datetime):
        return val.date()
    if isinstance(val, datetime.date):
        return val
    if isinstance(val, str):
        for fmt in ("%d.%m.%Y", "%Y-%m-%d"):
            try:
                return datetime.datetime.strptime(val.strip(), fmt).date()
            except ValueError:
                continue
    return None


def _parse_time_s3(val):
    """Parses time from Strona 3 E column value (datetime, time, float fraction, or string)."""
    if isinstance(val, datetime.datetime):
        return val.time()
    if isinstance(val, datetime.time):
        return val
    if isinstance(val, (int, float)):
        # Excel stores time as fraction of day (e.g. 0.375 = 09:00:00)
        total_sec = int(round(val * 86400))
        total_sec = total_sec % 86400
        h = total_sec // 3600
        m = (total_sec % 3600) // 60
        s = total_sec % 60
        return datetime.time(h, m, s)
    if isinstance(val, str):
        for fmt in ("%H:%M:%S", "%H:%M"):
            try:
                return datetime.datetime.strptime(val.strip(), fmt).time()
            except ValueError:
                continue
    return None


def _oblicz_warunki_srodowiskowe(app, wb, dane_zakladek, nr_pom, model, _cache_fg=None):
    """
    Uses dates/times from dane_zakladek (read from the PROTOCOL's Strona 3) to look
    up T/H from the environmental sensor file, writes T/H into F/G cells of the
    COPY's Strona 3 (wb), then uses Wzory.xls (already open in app) to compute
    min/max ranges.
    Returns: {"temp_min":..., "temp_max":..., "wilg_min":..., "wilg_max":...} or None.
    Wzory.xls is reverted to its original state after calculation (not saved).
    """
    # Warunki srodowiskowe (pomieszczenie) sa IDENTYCZNE dla wszystkich kopii — zaleza od
    # daty/godziny punktu, a nie od przyrzadu. Gdy komplet F/G jest juz w cache z pierwszej
    # kopii, nie ma po co otwierac protokolu ani pliku czujnika (16 MB) po raz kolejny.
    _bloki = [zd.get("block_idx", 0) for zd in dane_zakladek]
    _komplet_z_cache = bool(_cache_fg) and all(b in _cache_fg for b in _bloki)
    if _komplet_z_cache:
        print(f"    [Srodowisko] Komplet F/G z cache ({len(_bloki)} blokow) — "
              f"nie otwieram protokolu ani pliku czujnika.")

    ws_s3 = None
    if ARKUSZ_STRONA3 in {s.name for s in wb.sheets}:
        ws_s3 = wb.sheets[ARKUSZ_STRONA3]

    # Kopia pochodzi z szablonu bez Strona 3 — pisz F/G bezposrednio do protokolu.
    wb_protokol_s3 = None
    if ws_s3 is None and not _komplet_z_cache:
        protokol_sciezka = os.path.join(FOLDER, PROTOKOL_PLIK)
        try:
            for bk in app.books:
                if os.path.normcase(bk.fullname) == os.path.normcase(protokol_sciezka):
                    wb_protokol_s3 = bk
                    break
            if wb_protokol_s3 is None and os.path.exists(protokol_sciezka):
                wb_protokol_s3 = _open_book_hidden(app, protokol_sciezka, update_links=False)
            if wb_protokol_s3 is not None and ARKUSZ_STRONA3 in {s.name for s in wb_protokol_s3.sheets}:
                ws_s3 = wb_protokol_s3.sheets[ARKUSZ_STRONA3]
                print(f"    [Srodowisko] Brak Strona 3 w kopii — F/G zapisywane do protokolu.")
        except Exception as _e:
            _warn(f"Nie mozna otworzyc protokolu do zapisu F/G: {protokol_sciezka}\n"
                  f"        {type(_e).__name__}: {_e}", indent="    ")

    cache_czujnika = {}

    def _dane_miesiaca(rok, miesiac):
        key = (rok, miesiac)
        if key not in cache_czujnika:
            sciezka = _sciezka_pliku_czujnika(nr_pom, model, rok)
            cache_czujnika[key] = _wczytaj_dane_czujnika_miesiac(sciezka, rok, miesiac)
        return cache_czujnika[key]

    dane_th = []
    protokol_changed = False
    print(f"    [Srodowisko] Zakladki do przetworzenia: {len(dane_zakladek)}")
    for zd in dane_zakladek:
        data_val = zd.get("K4_val")
        czas_st_raw = zd.get("czas_start")
        czas_kon_raw = zd.get("czas_koniec")
        blk = zd.get("block_idx", "?")
        if data_val is None:
            print(f"      blok {blk}: brak daty, pomijam")
            continue
        data = _parse_date_s3(data_val)
        if data is None:
            print(f"      blok {blk}: nie mozna sparsowac daty, pomijam")
            continue
        block_idx = zd.get("block_idx", 0)
        row_base = START_ROW_S3 + block_idx * BLOK_S3

        # Jesli F/G juz wpisane w protokole — czytaj z nich i pomijaj wyszukiwanie czujnika.
        # Dane F/G sa identyczne dla wszystkich kopii — po pierwszym odczycie uzywamy cache.
        if _cache_fg is not None and block_idx in _cache_fg:
            f2, g2, f3, g3 = _cache_fg[block_idx]
            for t, h in [(f2, g2), (f3, g3)]:
                if t is not None:
                    dane_th.append((t, h))
            continue

        if ws_s3 is not None:
            f2 = ws_s3.range(f"F{row_base + 2}").value
            f3 = ws_s3.range(f"F{row_base + 3}").value
            g2 = ws_s3.range(f"G{row_base + 2}").value
            g3 = ws_s3.range(f"G{row_base + 3}").value
            if f2 is not None and f3 is not None:
                if _cache_fg is not None:
                    _cache_fg[block_idx] = (f2, g2, f3, g3)
                print(f"      blok {blk}: F/G juz wypelnione ({f2}/{g2}, {f3}/{g3}), pomijam czujnik")
                for t, h in [(f2, g2), (f3, g3)]:
                    if t is not None:
                        dane_th.append((t, h))
                continue

        rok_bloku = data.year
        mies_bloku = data.month
        dane = _dane_miesiaca(rok_bloku, mies_bloku)
        print(f"      blok {blk}: {data}  start={czas_st_raw!r}  koniec={czas_kon_raw!r}")
        print(f"        -> rekordow czujnika na {mies_bloku:02d}/{rok_bloku}: {len(dane)}")
        for czas_val, row_offset, etykieta in [
            (czas_st_raw, 2, "start"),
            (czas_kon_raw, 3, "koniec"),
        ]:
            czas = _parse_time_s3(czas_val)
            if czas is None:
                print(f"        -> {etykieta}: nie mozna sparsowac czasu ({czas_val!r}), pomijam")
                continue
            target_dt = datetime.datetime.combine(data, czas)
            temp, hum, odch_s = _szukaj_th_w_danych_czujnika(dane, target_dt)
            if temp is None and odch_s is not None:
                print(f"        -> {etykieta} {target_dt}: BRAK danych czujnika "
                      f"(najblizszy rekord {odch_s/60:.0f} min od punktu, limit "
                      f"{TOLERANCJA_CZUJNIKA_MIN} min) — F/G zostaja puste")
            else:
                print(f"        -> {etykieta} {target_dt}: T={temp}, H={hum}"
                      f"{f' (rekord {odch_s/60:.0f} min od punktu)' if odch_s else ''}")
            if ws_s3 is not None and temp is not None:
                ws_s3.range(f"F{row_base + row_offset}").value = temp
                protokol_changed = True
            if ws_s3 is not None and hum is not None:
                ws_s3.range(f"G{row_base + row_offset}").value = hum
            if temp is not None:
                dane_th.append((temp, hum))

    # Zapisz protokol tylko gdy cos zmieniono.
    if wb_protokol_s3 is not None and protokol_changed:
        try:
            wb_protokol_s3.save()
            print(f"    [Srodowisko] Zapisano F/G do Strona 3 protokolu.")
        except Exception as _e:
            print(f"    [UWAGA] Nie mozna zapisac protokolu po wpisaniu F/G: {_e}")

    if not dane_th:
        print("    [UWAGA] Brak danych T/H ze srodowiska — pomijam obliczenie zakresow.")
        return None

    # Find Wzory.xls in the currently open books
    wzory_wb = None
    for bk in app.books:
        if bk.name.lower() == "wzory.xls":
            wzory_wb = bk
            break
    if wzory_wb is None:
        print("    [UWAGA] Plik Wzory.xls nie jest otwarty — pomijam obliczenie zakresow srodowiskowych.")
        return None

    try:
        ws_env = wzory_wb.sheets["Środowiskowe"]
    except Exception:
        print("    [UWAGA] Brak arkusza 'Srodowiskowe' w Wzory.xls.")
        return None

    # Save originals to revert after calculation
    g116_orig = ws_env.range("G116").value
    g117_orig = ws_env.range("G117").value
    b_orig = ws_env.range("B115:B217").value
    c_orig = ws_env.range("C115:C217").value

    zakresy = None
    try:
        ws_env.range("G116").value = nr_pom
        ws_env.range("G117").value = model
        ws_env.range("B115:B217").clear_contents()
        ws_env.range("C115:C217").clear_contents()
        dane_th_valid = [(t, h) for t, h in dane_th if t is not None]
        for i, (temp, hum) in enumerate(dane_th_valid[:103]):
            ws_env.range(f"B{115 + i}").value = temp
            if hum is not None:
                ws_env.range(f"C{115 + i}").value = hum
        app.api.DisplayAlerts = False
        app.api.AskToUpdateLinks = False
        app.api.EnableEvents = False
        for _ws in wb.sheets:
            _ws.api.Calculate()
        app.api.EnableEvents = True
        temp_min = ws_env.range("L137").value
        temp_max = ws_env.range("N137").value
        wilg_min = ws_env.range("L138").value
        wilg_max = ws_env.range("N138").value
        zakresy = {
            "temp_min": temp_min,
            "temp_max": temp_max,
            "wilg_min": wilg_min,
            "wilg_max": wilg_max,
        }
    except Exception as e:
        print(f"    [UWAGA] Blad obliczen srodowiskowych: {e}")
    finally:
        try:
            ws_env.range("G116").value = g116_orig
            ws_env.range("G117").value = g117_orig
            ws_env.range("B115:B217").clear_contents()
            ws_env.range("C115:C217").clear_contents()
            if b_orig:
                ws_env.range("B115:B217").value = b_orig
            if c_orig:
                ws_env.range("C115:C217").value = c_orig
        except Exception:
            pass

    return zakresy


def _formatuj_zakres_srodowiskowy(v, miejsca=1):
    """
    Formatuje wartosc warunkow srodowiskowych do swiadectwa Word.

    miejsca=1 — temperatura otoczenia, np. '21,8' (przecinek dziesietny PL).
    miejsca=0 — wilgotnosc wzgledna, np. '30' — BEZ przecinka.

    Wilgotnosc w swiadectwach podaje sie w pelnych procentach; wczesniej kazda
    wartosc dostawala doklejone ',0' ('30,0 ÷ 54,0 %'), co nie odpowiada
    rozdzielczosci tego pomiaru.

    Zaokraglamy „w gore od polowy" (30,5 -> 31), a nie bankowo jak wbudowane
    round(), ktore dla 30,5 dalo by 30, a dla 31,5 — 32.
    """
    if v is None:
        return "—"
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    if miejsca == 0:
        return str(int(math.floor(f + 0.5)) if f >= 0 else -int(math.floor(-f + 0.5)))
    return f"{f:.{miejsca}f}".replace(".", ",")


def _oblicz_zakresy_srodowiskowe_z_istniejacych_kopii(folder, nazwy, dane_zakladek_per_kopia):
    """
    Computes per-copy environmental ranges for already-existing Excel copies
    (GENERUJ_EXCEL=False path). Opens each copy via xlwings, fills Strona 3
    F/G cells from sensor data, computes min/max via Wzory.xls, saves copy.
    Returns list of dicts (or None entries on failure).
    """
    n = len(nazwy)
    zakresy_per_kopia = []
    app = xw.App(visible=False, add_book=False)
    app.api.AutomationSecurity = 1
    app.api.DisplayAlerts = False
    app.api.AskToUpdateLinks = False
    app.api.Visible = False
    # Blokuj Workbook_Open/Worksheet_Calculate makr JUZ PRZED otwarciem plikow
    # linkowanych (Wzory.xls / Obliczenia tdp, RH, C.xls) — ich makra siegaja
    # \\plum4 i moga uruchamiac Worda, co zawiesza skrypt. UDF-y dzialaja i tak
    # (EnableEvents nie dotyczy funkcji, tylko zdarzen). Poszczegolne operacje
    # przywracaja EnableEvents wg potrzeby.
    app.api.EnableEvents = False
    try:
        app.api.AutoRecover.Enabled = False  # nie twórz plików autoodzyskiwania (.xar)
    except Exception:
        pass
    try:
        for plik_link in PLIKI_LINKOWANE:
            sciezka_link = os.path.join(folder, plik_link)
            if os.path.exists(sciezka_link):
                try:
                    _open_book_hidden(app, sciezka_link, update_links=False)
                    _ok(f"Otwarto plik linkowany: {plik_link}")
                except Exception as exc:
                    _warn(f"Nie mozna otworzyc pliku linkowanego: {plik_link}\n"
                          f"      Sciezka lokalna: {sciezka_link}\n"
                          f"      Blad: {type(exc).__name__}: {exc}")
        for j, nazwa in enumerate(nazwy):
            print(f"    [Srodowisko {j+1:>{len(str(n))}}/{n}] {nazwa}")
            sciezka = os.path.join(folder, nazwa)
            dane_zak = dane_zakladek_per_kopia[j] if j < len(dane_zakladek_per_kopia) else []
            wb = _open_book_hidden(app, sciezka, update_links=False)
            zakresy = None
            try:
                zakresy = _oblicz_warunki_srodowiskowe(app, wb, dane_zak, NR_POMIESZCZENIA, MODEL_CZUJNIKA)
                if zakresy is not None:
                    app.api.DisplayAlerts = False
                    app.api.AskToUpdateLinks = False
                    app.api.EnableEvents = False
                    for _ws in wb.sheets:
                        _ws.api.Calculate()
                    app.api.EnableEvents = True
                    wb.save()
            except Exception as e:
                print(f"    [UWAGA] Blad obliczen srodowiskowych dla '{nazwa}': {e}")
            finally:
                try:
                    wb.close()
                except Exception:
                    pass
            zakresy_per_kopia.append(zakresy)
    finally:
        app.quit()
    return zakresy_per_kopia


# =============================================================================
# FUNKCJE POMOCNICZE
# =============================================================================

def _cell_to_str(value):
    """
    Konwertuje wartość komórki do stringa.
    - None          → ""
    - float bez ułamka (np. 20.0) → "20"
    - float z ułamkiem (np. 20.3) → "20.3"
    - pozostałe     → str(value).strip()
    """
    if value is None:
        return ""
    if isinstance(value, float):
        return str(int(value)) if value == int(value) else str(value)
    return str(value).strip()


def _excel_color_to_hex(ole_color):
    """Konwertuje kolor COM Excela (OLE BGR int) do postaci #RRGGBB."""
    if ole_color in (None, ""):
        return None
    try:
        kolor = int(ole_color)
    except (TypeError, ValueError):
        return None

    r = kolor & 0xFF
    g = (kolor >> 8) & 0xFF
    b = (kolor >> 16) & 0xFF
    return f"#{r:02X}{g:02X}{b:02X}"


def _kolor_komorki_hex_xlwings(cell):
    """Zwraca kolor komorki jako #RRGGBB (DisplayFormat->Interior fallback)."""
    try:
        ole_color = cell.api.DisplayFormat.Interior.Color
    except Exception:
        try:
            ole_color = cell.api.Interior.Color
        except Exception:
            ole_color = None
    return _excel_color_to_hex(ole_color)


def _czy_bierz_dane_po_kolorze(hex_color):
    """Decyduje czy pobierac wartosc komorki na podstawie koloru."""
    if not FILTRUJ_KOLOR_S3:
        return True

    if hex_color is None:
        return BIERZ_INNE_KOLORY_S3

    kolor = str(hex_color).upper()
    if kolor == KOLOR_AKTYWNY_S3.upper():
        return True
    if kolor == KOLOR_POMIJANY_S3.upper():
        return False
    return BIERZ_INNE_KOLORY_S3


def _wartosc_s3_po_kolorze_xlwings(cell):
    """Zwraca wartosc komorki tylko gdy kolor jest dozwolony przez konfiguracje."""
    if _czy_bierz_dane_po_kolorze(_kolor_komorki_hex_xlwings(cell)):
        return cell.value
    return None


def _wartosc_z_scalonej_komorki_xlwings(cell):
    """Zwraca wartosc z obszaru scalonego (top-left), fallback: zwykla wartosc."""
    try:
        if cell.api.MergeCells:
            return cell.api.MergeArea.Cells(1, 1).Value
    except Exception:
        pass
    return cell.value


def _czy_wartosc_niepusta(value):
    """True dla wartosci, ktore traktujemy jako realne dane pomiarowe."""
    if value is None:
        return False
    if isinstance(value, str) and value.strip() == "":
        return False
    return True


def _get_number_format(cell):
    """Zwraca NumberFormat komorki xlwings; None dla General i @."""
    try:
        fmt = cell.api.NumberFormat
        if fmt and fmt.upper() not in ('GENERAL', '@'):
            return fmt
    except Exception:
        pass
    return None


def _czy_blok_ef_aktywny(ef_blok):
    """Sprawdza czy blok ma co najmniej jedna niepusta wartosc w E/F."""
    if not isinstance(ef_blok, dict):
        return False

    for key in ("E_dane", "F_dane"):
        for value in (ef_blok.get(key) or []):
            if _czy_wartosc_niepusta(value):
                return True
    return False


def _wybierz_aktywne_bloki_kopii(dane_zakladek, dane_ef_kopia):
    """
    Zwraca aktywne bloki dla jednej kopii.
    Aktywny blok = ma przynajmniej jedna wartosc E/F po filtracji kolorow.
    """
    aktywne_idx = [i for i, ef in enumerate(dane_ef_kopia) if _czy_blok_ef_aktywny(ef)]
    dane_zakladek_kopia = [dane_zakladek[i] for i in aktywne_idx if i < len(dane_zakladek)]
    dane_ef_kopia_aktywne = [dane_ef_kopia[i] for i in aktywne_idx]
    return aktywne_idx, dane_zakladek_kopia, dane_ef_kopia_aktywne


def _odfiltruj_przyrzady_bez_danych(dane_s2, dane_ef, f24_per_kopia):
    """
    Usuwa z obiegu przyrzady, ktore nie maja ANI JEDNEGO aktywnego bloku pomiarow.

    Lista kopii do zrobienia pochodzi ze Strony 2 (tabela przyrzadow), a nie z
    kolorow na Stronie 3. Gdy uzytkownik wyszarzy wszystkie pomiary przyrzadu,
    zeby go pominac, przyrzad nadal siedzi na Stronie 2 — bez tego filtra
    powstawala pusta kopia (sam arkusz Wyniki) i do tego bezuzyteczne
    swiadectwo Word z zerowa tabela kalibracji.

    Numer przyrzadu z protokolu jest zapamietywany w rekordzie, zeby nazwa
    pozostalej kopii dalej zgadzala sie z pozycja na Stronie 2 (np. '... - 4 -').

    Zwraca (dane_s2, dane_ef, f24_per_kopia, pominiete) po filtracji.
    """
    zachowane_s2, zachowane_ef, zachowane_f24, pominiete = [], [], [], []
    for j, rekord in enumerate(dane_s2):
        rekord["_nr_przyrzadu"] = j + 1          # pozycja w tabeli Strona 2
        ef_kopia = dane_ef[j] if j < len(dane_ef) else []
        ma_dane = any(_czy_blok_ef_aktywny(ef) for ef in ef_kopia)
        if ma_dane:
            zachowane_s2.append(rekord)
            zachowane_ef.append(ef_kopia)
            zachowane_f24.append(f24_per_kopia[j] if j < len(f24_per_kopia) else None)
        else:
            pominiete.append((j + 1, rekord))
    return zachowane_s2, zachowane_ef, zachowane_f24, pominiete


def _zbuduj_dane_zakladek_per_kopia(dane_zakladek, dane_ef):
    """Buduje liste zakladek roboczych osobno dla kazdej kopii."""
    wynik = []
    for dane_ef_kopia_surowe in dane_ef:
        if USUWAJ_PUSTE_BLOKI_KOPII_S3:
            _, dane_zakladek_kopia, _ = _wybierz_aktywne_bloki_kopii(
                dane_zakladek,
                dane_ef_kopia_surowe,
            )
        else:
            dane_zakladek_kopia = list(dane_zakladek)
        wynik.append(dane_zakladek_kopia)
    return wynik


def _formatuj_zakresy_wierszy_s3(indeksy, start_row, blok):
    """Formatuje indeksy blokow do postaci zakresow wierszy, np. 20-24."""
    zakresy = []
    for i in indeksy:
        r0 = start_row + i * blok
        zakresy.append(f"{r0}-{r0 + blok - 1}")
    return ", ".join(zakresy)



def _nazwa_pliku_z_linku(target):
    """
    Wyciaga nazwe pliku (basename) z dowolnej postaci sciezki/linku zewnetrznego
    Excela ('\\\\serwer\\...', 'C:\\...', 'file:///C:/...', wzgledna nazwa).
    Excel zapisuje spacje w Target jako %20 (a przecinki bez kodowania) —
    bez unquote() basename z plikow ze spacjami w nazwie ('Obliczenia tdp, RH, C.xls')
    nigdy nie trafi w klucz slownika linkow, mimo ze plik jest ten sam.
    """
    if not isinstance(target, str):
        return None
    baza = os.path.basename(target.replace("\\", "/"))
    try:
        baza = urllib.parse.unquote(baza)
    except Exception:
        pass
    return baza


def _przywroc_linki_w_xml(sciezka_pliku, linki_serwerowe, cicho=False):
    """
    Zmienia sciezki linkow zewnetrznych bezposrednio w XML pliku xlsx
    (bez otwierania przez Excel), dzieki czemu zakeszowane wartosci formul
    pozostaja nienaruszone i sa widoczne po otwarciu pliku bez aktualizacji linkow.

    Uzywane dwukierunkowo:
      - na koncu: sciezki serwerowe (UNC plum4) — przywrocenie do publikacji,
      - PRZED otwarciem kopii: sciezki lokalne — zeby Excel NIE szukal serwera
        (niedostepny serwer => zawieszenie SMB => RPC crash przy Open).

    Dlaczego nie uzywamy ChangeLink przez xlwings:
    ChangeLink() wymusza przeliczenie formul. Jesli nowa sciezka UNC jest
    niedostepna z maszyny ze skryptem, Excel zapisuje bledy jako cache —
    i inne osoby widza puste komorki po kliknieciu 'Nie' w dialogu linkow.
    cicho=True wycisza logowanie (przy masowym przepisywaniu przed otwarciem).
    """
    if not linki_serwerowe:
        return

    linki_po_nazwie = {k.lower(): v for k, v in linki_serwerowe.items()}
    sciezka_tmp = sciezka_pliku + "._xltmp_"
    zmienione_pliki = 0
    zmiany_log = []  # (stara_sciezka, nowa_sciezka)

    try:
        with zipfile.ZipFile(sciezka_pliku, 'r') as zin, \
             zipfile.ZipFile(sciezka_tmp, 'w', compression=zipfile.ZIP_DEFLATED) as zout:

            for item in zin.infolist():
                data = zin.read(item.filename)

                if (item.filename.startswith('xl/externalLinks/_rels/')
                        and item.filename.endswith('.rels')):
                    tekst = data.decode('utf-8')
                    oryg = tekst

                    def _podmien(m):
                        target = m.group(1)
                        t_basename = _nazwa_pliku_z_linku(target)
                        nowa = linki_po_nazwie.get(t_basename.lower()) if t_basename else None
                        if nowa:
                            zmiany_log.append((target, nowa))
                            return f'Target="{nowa}"'
                        return m.group(0)

                    tekst = re.sub(r'Target="([^"]+)"', _podmien, tekst)
                    if tekst != oryg:
                        zmienione_pliki += 1
                    data = tekst.encode('utf-8')

                zout.writestr(item, data)

        os.replace(sciezka_tmp, sciezka_pliku)
        if not cicho:
            if zmienione_pliki:
                _ok(f"Przywrocono linki zewnetrzne ({zmienione_pliki} rel) w: "
                    f"{os.path.basename(sciezka_pliku)}", indent="    ")
                for stara, nowa in zmiany_log:
                    print(f"        stara: {stara}")
                    print(f"        nowa:  {nowa}")
            else:
                print(f"    [XML] Brak linkow do zastapienia w: {os.path.basename(sciezka_pliku)}")
    except Exception as exc:
        if os.path.exists(sciezka_tmp):
            try:
                os.remove(sciezka_tmp)
            except Exception:
                pass
        _warn(f"Nie mozna zmienic linkow XML w '{os.path.basename(sciezka_pliku)}'\n"
              f"      {type(exc).__name__}: {exc}", indent="    ")


def _odczytaj_kalibracje_dla_istniejacych_kopii(folder, nazwy_kopii, ark_wyniki):
    """Czyta kalibracje z juz istniejacych kopii Excel (bez etapu tworzenia kopii)."""
    n = len(nazwy_kopii)
    dane_kalibracji = []
    klasa_wilg_per_kopia = []

    app = xw.App(visible=False, add_book=False)
    app.api.AutomationSecurity = 1
    app.api.DisplayAlerts = False
    app.api.AskToUpdateLinks = False
    app.api.Visible = False
    # Blokuj Workbook_Open/Worksheet_Calculate makr JUZ PRZED otwarciem plikow
    # linkowanych (Wzory.xls / Obliczenia tdp, RH, C.xls) — ich makra siegaja
    # \\plum4 i moga uruchamiac Worda, co zawiesza skrypt. UDF-y dzialaja i tak
    # (EnableEvents nie dotyczy funkcji, tylko zdarzen). Poszczegolne operacje
    # przywracaja EnableEvents wg potrzeby.
    app.api.EnableEvents = False
    try:
        app.api.AutoRecover.Enabled = False  # nie twórz plików autoodzyskiwania (.xar)
    except Exception:
        pass
    try:
        linked_wbs = []
        sciezki_linkowane = {}
        for plik_link in PLIKI_LINKOWANE:
            sciezka_link = os.path.join(folder, plik_link)
            if os.path.exists(sciezka_link):
                try:
                    lwb = _open_book_hidden(app, sciezka_link, update_links=False)
                    linked_wbs.append(lwb)
                    sciezki_linkowane[plik_link.lower()] = lwb.fullname
                    _ok(f"Otwarto plik linkowany: {plik_link}")
                except Exception as exc:
                    _warn(f"Nie mozna otworzyc pliku linkowanego: {plik_link}\n"
                          f"      Sciezka lokalna: {sciezka_link}\n"
                          f"      Blad: {type(exc).__name__}: {exc}")
            else:
                _warn(f"Brak pliku linkowanego w folderze roboczym: {plik_link}\n"
                      f"      Sciezka serwerowa: {_info_serwer(plik_link)}\n"
                      f"      Oczekiwano w: {sciezka_link}")

        for j, nazwa in enumerate(nazwy_kopii, start=1):
            sciezka = os.path.join(folder, nazwa)
            print(f"    [Kalibracja {j:>{len(str(n))}}/{n}] {nazwa}")
            kal, klasa = _odczytaj_kalibracje_xlwings(
                app, sciezka, ark_wyniki, enable_diag=(j == 1), sciezki_linkowane=sciezki_linkowane)
            dane_kalibracji.append(kal or [])
            klasa_wilg_per_kopia.append(klasa)

        for lwb in linked_wbs:
            try:
                lwb.close()
            except Exception:
                pass
    finally:
        app.quit()

    return dane_kalibracji, klasa_wilg_per_kopia


def _czy_protokol_cc04(sciezka):
    """Wykrywa nowy typ protokolu na podstawie nazwy pliku (CC-04)."""
    nazwa = os.path.basename(str(sciezka or ""))
    return "CC-04" in nazwa.upper()


def _wykryj_tag_cc04(wartosc):
    """Wykrywa tag typu CC-04: LG/LD/PD/PG."""
    if wartosc is None:
        return None

    tekst = str(wartosc).strip().upper()
    if not tekst:
        return None

    compact = "".join(ch for ch in tekst if ch.isalnum())
    dozwolone = ("LG", "LD", "PD", "PG")

    if compact in dozwolone:
        return compact

    if "CC04" in compact:
        tail = compact.split("CC04", 1)[1]
        for tag in dozwolone:
            if tail.startswith(tag):
                return tag

    for tag in dozwolone:
        if compact.endswith(tag) and len(compact) <= 8:
            return tag

    return None


def _parametry_typu_cc04(rekord):
    """Zwraca mapowanie komorek K11/K12/K13/K17 dla kopii CC-04."""
    if not isinstance(rekord, dict):
        return None

    tag = rekord.get("CC04_TAG")
    if not tag:
        tag = _wykryj_tag_cc04(rekord.get("CC04_RAW"))
    if not tag:
        return None

    return MAPOWANIE_TYPU_CC04.get(str(tag).upper())


def _round_half_away_from_zero(value):
    """Zaokragla do calkowitych jak w Excel ROUND(..., 0)."""
    if value >= 0:
        return int(math.floor(value + 0.5))
    return int(math.ceil(value - 0.5))


def _cell_to_sheet_name_part(value):
    """
    Konwertuje wartosc do fragmentu nazwy zakladki i zaokragla liczby do calosci.
    Przyklad: 24.9 -> "25", 48.5 -> "49", -19.9 -> "-20".
    """
    if value is None:
        return ""

    if isinstance(value, bool):
        return str(value)

    if isinstance(value, (int, float)):
        return str(_round_half_away_from_zero(float(value)))

    text = str(value).strip()
    if not text:
        return ""

    # Dla tekstow numerycznych (np. "24,9" lub "24.9")
    probe = text.replace(",", ".")
    try:
        num = float(probe)
    except ValueError:
        return text
    return str(_round_half_away_from_zero(num))


_EXCEL_NIEDOZWOLONE_ZNAKI = set(r'[]:*?/\\')


def _unikalne_nazwy_zakladek(dane_zakladek):
    """
    Buduje unikalne nazwy arkuszy zgodne z ograniczeniami Excela:
    - max 31 znakow
    - bez: []:*?/\\
    - bez duplikatow (case-insensitive)
    """
    wynik = []
    zajete = set()

    for i, zd in enumerate(dane_zakladek, start=1):
        raw = _cell_to_str((zd or {}).get("nazwa"))
        base = "".join("_" if ch in _EXCEL_NIEDOZWOLONE_ZNAKI else ch for ch in raw)
        base = base.strip().strip("'")
        if not base:
            base = f"Zakladka {i}"
        base = base[:31]

        candidate = base
        licznik = 2
        while candidate.lower() in zajete:
            suffix = f" ({licznik})"
            max_len = max(1, 31 - len(suffix))
            candidate = base[:max_len].rstrip() + suffix
            licznik += 1

        zajete.add(candidate.lower())
        wynik.append(candidate)

    return wynik


MIESIACE_GEN = {
    1: "stycznia",  2: "lutego",       3: "marca",     4: "kwietnia",
    5: "maja",      6: "czerwca",      7: "lipca",     8: "sierpnia",
    9: "września", 10: "października", 11: "listopada", 12: "grudnia",
}


# =============================================================================
# WCZYTYWANIE DANYCH Z PROTOKOŁU
# =============================================================================

def wczytaj_dane_z_protokolu_s2(sciezka, arkusz, start_row):
    """
    Czyta listę kopii do wygenerowania z Strona 2.

    Kolumna A (1)  – warunek zatrzymania (pusta = koniec)
    Kolumna E (5)  – wartość zastępująca 'RH (CC)' w nazwie kopii
    Kolumna O (15) – wartość zastępująca 'xxx' w nazwie kopii

    Zwraca: [{"O": "133", "E": "1010223"}, ...]
    """
    wb = openpyxl.load_workbook(sciezka, read_only=True, data_only=True)
    if arkusz not in wb.sheetnames:
        wb.close()
        raise ValueError(f"Brak arkusza '{arkusz}' w pliku '{sciezka}'")

    ws = wb[arkusz]
    dane = []
    row = start_row

    while True:
        val_a = ws.cell(row=row, column=1).value
        if val_a is None or str(val_a).strip() == "":
            break

        val_O = ws.cell(row=row, column=15).value   # kolumna O
        val_E = ws.cell(row=row, column=5).value    # kolumna E

        if val_O is None or val_E is None:
            print(f"  [OSTRZEŻENIE] Wiersz {row} w '{arkusz}': brak wartości O lub E — pomijam wiersz.")
            row += 1
            continue

        dane.append({"O": _cell_to_str(val_O), "E": _cell_to_str(val_E)})
        row += 1

    wb.close()
    return dane


def wczytaj_zakladki_z_protokolu_s3(sciezka, arkusz, start_row, blok):
    """
    Czyta definicje zakładek z Strona 3 w blokach po 5 wierszy.

    Dla każdego bloku (startowy wiersz r0):
      - Kolumna A (1) w r0 – warunek zatrzymania (pusta = koniec)
      - Kolumna B (2) w r0 – pierwsza część nazwy zakładki
      - Kolumna C (3) w r0 – druga część nazwy zakładki  → nazwa = "{B}, {C}"
      - Kolumna L (12), wiersze r0..r0+4 – dane dla C15:C19
      - Kolumna M (13), wiersze r0..r0+4 – dane dla D15:D19

    Zwraca: [{"nazwa": "20.3, -", "L_dane": [5 val], "M_dane": [5 val]}, ...]
    """
    wb = openpyxl.load_workbook(sciezka, read_only=True, data_only=True)
    if arkusz not in wb.sheetnames:
        wb.close()
        raise ValueError(f"Brak arkusza '{arkusz}' w pliku '{sciezka}'")

    ws = wb[arkusz]
    wynik = []
    blok_idx = 0

    while True:
        r0 = start_row + blok_idx * blok
        val_a = ws.cell(row=r0, column=1).value
        if val_a is None or str(val_a).strip() == "":
            break

        val_b = _cell_to_sheet_name_part(ws.cell(row=r0, column=2).value)
        val_c = _cell_to_sheet_name_part(ws.cell(row=r0, column=3).value)
        nazwa = f"{val_b}, {val_c}"

        L_dane = [ws.cell(row=r0 + k, column=12).value for k in range(blok)]
        M_dane = [ws.cell(row=r0 + k, column=13).value for k in range(blok)]

        wynik.append({"nazwa": nazwa, "L_dane": L_dane, "M_dane": M_dane})
        blok_idx += 1

    wb.close()
    return wynik


def wczytaj_dane_ef_z_protokolu_s3(sciezka, arkusz, n_kopii, n_zakladek,
                                    start_row, blok,
                                    start_col_e, start_col_f, krok):
    """
    Czyta dane E/F z Strona 3 dla każdej kopii i każdej zakładki.

    Kopia j (0-bazowany):
      kolumna E → start_col_e + j * krok   (Q dla j=0, S dla j=1, U dla j=2, …)
      kolumna F → start_col_f + j * krok   (R dla j=0, T dla j=1, V dla j=2, …)
    Zakładka i (0-bazowany):
      wiersze → start_row + i * blok  ..  start_row + i * blok + (blok-1)

    Zwraca: dane_ef[j][i] = {"E_dane": [5 val], "F_dane": [5 val]}
    """
    wb = openpyxl.load_workbook(sciezka, read_only=True, data_only=True)
    if arkusz not in wb.sheetnames:
        wb.close()
        raise ValueError(f"Brak arkusza '{arkusz}' w pliku '{sciezka}'")

    ws = wb[arkusz]
    dane_ef = []

    for j in range(n_kopii):
        col_e = start_col_e + j * krok
        col_f = start_col_f + j * krok
        zakl = []
        for i in range(n_zakladek):
            r0 = start_row + i * blok
            zakl.append({
                "E_dane": [ws.cell(row=r0 + k, column=col_e).value for k in range(blok)],
                "F_dane": [ws.cell(row=r0 + k, column=col_f).value for k in range(blok)],
            })
        dane_ef.append(zakl)

    wb.close()
    return dane_ef


def wczytaj_wszystko_xlwings(sciezka, ark_s2, ark_s3,
                              start_s2, start_s3, blok,
                              start_col_e, start_col_f, krok):
    """
    Otwiera plik protokołu RAZ przez xlwings (COM Excel) i odczytuje:
      - listę kopii z Strona 2
            - definicje zakładek + dane C/D z Strona 3
                (CC: L/M, CC-04: C z K/L/M/N wg typu i D z O)
      - dane E/F dla każdej kopii z Strona 3
    Używaj tej funkcji zamiast trzech osobnych wywołań opartych na openpyxl,
    gdy protokół zawiera formuły z niezakeszowanymi wartościami.
    Zwraca: (dane_s2, dane_zakladek, dane_ef)
    """
    app = xw.App(visible=False, add_book=False)
    app.api.AutomationSecurity = 1   # msoAutomationSecurityLow — wlacza makra bez pytania
    app.api.DisplayAlerts = False
    app.api.AskToUpdateLinks = False
    app.api.Visible = False
    # Blokuj Workbook_Open/Worksheet_Calculate makr JUZ PRZED otwarciem plikow
    # linkowanych (Wzory.xls / Obliczenia tdp, RH, C.xls) — ich makra siegaja
    # \\plum4 i moga uruchamiac Worda, co zawiesza skrypt. UDF-y dzialaja i tak
    # (EnableEvents nie dotyczy funkcji, tylko zdarzen). Poszczegolne operacje
    # przywracaja EnableEvents wg potrzeby.
    app.api.EnableEvents = False
    try:
        app.api.AutoRecover.Enabled = False  # nie twórz plików autoodzyskiwania (.xar)
    except Exception:
        pass
    try:
        wb = _open_book_hidden(app, sciezka, update_links=False)

        protokol_cc04 = _czy_protokol_cc04(sciezka)
        przesuniecie_cc04 = PRZESUNIECIE_STARTU_KOL_CC04 if protokol_cc04 else 0
        start_col_e_eff = start_col_e + przesuniecie_cc04
        start_col_f_eff = start_col_f + przesuniecie_cc04

        # --- Strona 2: lista kopii ---
        ws2 = wb.sheets[ark_s2]
        dane_s2 = []
        row = start_s2
        while True:
            val_a = ws2.cells(row, 1).value
            if val_a is None or str(val_a).strip() == "":
                break
            val_O = ws2.cells(row, 15).value   # kolumna O
            val_E = ws2.cells(row, 5).value    # kolumna E
            if val_O is None or val_E is None:
                print(f"  [OSTRZEŻENIE] Wiersz {row} w '{ark_s2}': brak O lub E — pomijam.")
                row += 1
                continue
            val_B = ws2.cells(row, 2).value   # kolumna B -> E5 kopii
            val_D = ws2.cells(row, 4).value   # kolumna D -> E6 kopii
            val_F = ws2.cells(row, 6).value   # kolumna F -> [nr_ewid] w Word
            val_K = ws2.cells(row, 11).value  # kolumna K -> H57 kopii
            dane_s2.append({
                "O": _cell_to_str(val_O),
                "E": _cell_to_str(val_E),
                "B": val_B,
                "D": val_D,
                "F": val_F,
                "K": val_K,
                "IS_CC04_PROTO": protokol_cc04,
                "CC04_RAW": "",
                "CC04_TAG": None,
            })
            row += 1

        # --- Strona 3: zakładki + dane L/M ---
        ws3 = wb.sheets[ark_s3]
        dane_zakladek = []
        blok_idx = 0
        while True:
            r0 = start_s3 + blok_idx * blok
            val_a = ws3.cells(r0, 1).value
            if val_a is None or str(val_a).strip() == "":
                break
            _raw_b = ws3.cells(r0, 2).value
            _raw_c = ws3.cells(r0, 3).value
            val_b = _cell_to_sheet_name_part(_raw_b)
            val_c = _cell_to_sheet_name_part(_raw_c)
            nazwa = f"{val_b}, {val_c}"
            # Surowe nastawy zostawiamy do dopasowania nazwy zakladki do punktu z PZ.
            _nast_t  = pz_dane._do_float(_raw_b)
            _nast_rh = pz_dane._do_float(_raw_c)
            # Dane L/M pobieramy tylko z komorek o dozwolonym kolorze.
            L_dane = []; L_fmt = []
            M_dane = []; M_fmt = []
            for k in range(blok):
                cl = ws3.cells(r0 + k, 12)
                L_dane.append(_wartosc_s3_po_kolorze_xlwings(cl))
                L_fmt.append(_get_number_format(cl))
                cm = ws3.cells(r0 + k, 13)
                M_dane.append(_wartosc_s3_po_kolorze_xlwings(cm))
                M_fmt.append(_get_number_format(cm))
            k4_val    = ws3.cells(r0 + 1, 5).value  # E(r0+1): data pomiaru -> K4 kopii
            czas_st   = ws3.cells(r0 + 2, 5).value  # E(r0+2): czas startu bloku
            czas_kon  = ws3.cells(r0 + 3, 5).value  # E(r0+3): czas konca bloku
            # O25: woda/szron — kol. P (16) dla CC, kol. R (18) dla CC-04
            _col_o25 = 18 if protokol_cc04 else 16
            _o25_raw = ws3.cells(r0, _col_o25).value
            _o25_str = str(_o25_raw).strip() if _o25_raw is not None else ""
            o25_val  = "N/A" if (_o25_str == "" or _o25_str == "-") else _o25_str
            dane_zakladek.append({
                "nazwa": nazwa,
                "block_idx": blok_idx,
                "L_dane": L_dane, "L_fmt": L_fmt,
                "M_dane": M_dane, "M_fmt": M_fmt,
                "K4_val": k4_val,
                "czas_start": czas_st,
                "czas_koniec": czas_kon,
                "o25_val": o25_val,
                "nast_t": _nast_t,      # surowa nastawa komory — do dopasowania punktu z PZ
                "nast_rh": _nast_rh,
            })
            blok_idx += 1

        # Nazwy zakladek wg punktow ZAMOWIONYCH w PZ (nastawa komory bywa inna niz nominal).
        _nazwy_zakladek_z_pz(dane_zakladek)

        # --- Strona 3: dane E/F per kopia ---
        n_kopii = len(dane_s2)
        n_zakladek = len(dane_zakladek)
        dane_ef = []

        if n_kopii > 0 and n_zakladek > 0:
            for j in range(n_kopii):
                col_e = start_col_e_eff + j * krok
                col_f = start_col_f_eff + j * krok
                col_c_cc04 = None
                col_d_cc04 = KOLUMNA_D_CC04_S3

                if protokol_cc04 and j < len(dane_s2):
                    raw_typ = _wartosc_z_scalonej_komorki_xlwings(ws3.cells(WIERSZ_TYPU_CC04_S3, col_e))
                    dane_s2[j]["CC04_RAW"] = _cell_to_str(raw_typ)
                    dane_s2[j]["CC04_TAG"] = _wykryj_tag_cc04(raw_typ)
                    tag = dane_s2[j]["CC04_TAG"]
                    if tag:
                        col_c_cc04 = MAPOWANIE_KOLUMNY_C_CC04.get(str(tag).upper())

                zakl = []
                for i in range(n_zakladek):
                    r0 = start_s3 + i * blok
                    e_dane = []; e_fmt = []
                    f_dane = []; f_fmt = []
                    for k in range(blok):
                        ce = ws3.cells(r0 + k, col_e)
                        e_dane.append(_wartosc_s3_po_kolorze_xlwings(ce))
                        e_fmt.append(_get_number_format(ce))
                        cf = ws3.cells(r0 + k, col_f)
                        f_dane.append(_wartosc_s3_po_kolorze_xlwings(cf))
                        f_fmt.append(_get_number_format(cf))

                    if protokol_cc04:
                        # Dla CC-04 C i D sa rowniez zalezne od kopii (po typie z wiersza 14).
                        # C/D bierzemy surowo z kolumn (bez filtrowania kolorem):
                        #   C <- K/L/M/N wg typu, D <- O.
                        c_dane = []; c_fmt = []
                        d_dane = []; d_fmt = []
                        col_src = col_c_cc04 if col_c_cc04 else 12
                        for k in range(blok):
                            cc = ws3.cells(r0 + k, col_src)
                            c_dane.append(cc.value)
                            c_fmt.append(_get_number_format(cc))
                            cd = ws3.cells(r0 + k, col_d_cc04)
                            d_dane.append(cd.value)
                            d_fmt.append(_get_number_format(cd))
                        zakl.append({
                            "C_dane": c_dane, "C_fmt": c_fmt,
                            "D_dane": d_dane, "D_fmt": d_fmt,
                            "E_dane": e_dane, "E_fmt": e_fmt,
                            "F_dane": f_dane, "F_fmt": f_fmt,
                        })
                    else:
                        zakl.append({
                            "E_dane": e_dane, "E_fmt": e_fmt,
                            "F_dane": f_dane, "F_fmt": f_fmt,
                        })
                dane_ef.append(zakl)

        # --- Strona 3: F24 dla arkusza Wyniki, per kopia (wiersz 17) ---
        # Wymaganie: czytaj wartosc ze scalonej komorki Q:R17 (lub S:T17 dla CC-04),
        # z kolejnymi kopiami co 2 kolumny w prawo.
        f24_per_kopia = []
        for j in range(n_kopii):
            col = start_col_e_eff + j * krok
            f24_per_kopia.append(_wartosc_z_scalonej_komorki_xlwings(ws3.cells(17, col)))

        wb.close()
        return dane_s2, dane_zakladek, dane_ef, f24_per_kopia
    finally:
        app.quit()


# =============================================================================
# GENEROWANIE NAZWY KOPII
# =============================================================================

def _numer_do_nazwy(rekord):
    """
    Numer przyrzadu uzywany w NAZWIE kopii: nr fabryczny (Strona 2 kol. E), a gdy go brak
    ('-' albo pusto) — nr ewidencyjny (kol. F).

    Czesc przyrzadow nie ma nadanego numeru fabrycznego i identyfikuje sie je numerem
    ewidencyjnym; bez tego zapasu nazwa konczyla sie samym mysleikiem ('... - 3 - -.xlsx')
    i pliki roznych przyrzadow bylyby nierozroznialne.
    """
    fabr = _cell_to_str(rekord.get("E")).strip()
    if fabr and fabr != "-":
        return fabr
    ewid = _cell_to_str(rekord.get("F")).strip()
    return ewid if ewid and ewid != "-" else fabr


def generuj_nazwe_pliku(szablon_nazwa, wartosc_O, wartosc_E, nr_przyrzadu=None):
    """
    Zastępuje w nazwie pliku szablonu:
      'xxx'      → wartosc_O        (numer zlecenia)
      '- N -'    → numer przyrzadu  (gdy podano nr_przyrzadu)
      'RH (CC)'  → wartosc_E        (nr fabryczny przyrzadu)

    Numer przyrzadu stoi w nazwie tuz PRZED nr fabrycznym ('... - 1 - RH (CC).xlsx')
    i musi odpowiadac pozycji przyrzadu w protokole (Przyrzady wzorcowane 1, 2, 3...).
    Wczesniej zostawala szablonowa '1' przy kazdej kopii, wiec pliki roznych przyrzadow
    mialy ten sam numer. Podmieniamy TYLKO liczbe przylegajaca do miejsca na nr fabryczny
    — inne liczby w nazwie (ILAJ 5.4_11#21, Wer.12, daty) zostaja nietkniete.
    """
    nazwa = szablon_nazwa
    if nr_przyrzadu is not None:
        nazwa = re.sub(r'-\s*\d+\s*-\s*(?=RH\s*\(CC\))',
                       f'- {nr_przyrzadu} - ', nazwa, count=1)
    nazwa = nazwa.replace("xxx", wartosc_O, 1)
    nazwa = nazwa.replace("RH (CC)", wartosc_E, 1)
    return nazwa


def _bezpieczna_nazwa_pliku(nazwa):
    """Zastępuje znaki niedozwolone w nazwach plików Windows znakiem '_'."""
    for ch in r'\/:*?"<>|':
        nazwa = nazwa.replace(ch, "_")
    return nazwa


def _parsuj_nazwe_pliku(nowa_nazwa):
    """
    Parsuje nazwe pliku kopii i zwraca (prefiks, rok, numer_fabryczny, typ).
    Przyklad: '133_LA_TH_2026 - ILAJ 5.4_11#21 - ... - 1010223.xlsx'
              -> ('133', '2026', '1010223', 'ILAJ 5.4_11#21')
    Prefiks         : czesc przed '_LA_TH_'
    Rok             : 4 cyfry po '_LA_TH_'
    Numer fabryczny : ostatni segment po ' - ' (bez rozszerzenia)
    Typ             : pierwszy segment po roku (model/typ przyrzadu pomiarowego)
    """
    nazwa = os.path.splitext(nowa_nazwa)[0]
    if "_LA_TH_" in nazwa:
        prefiks, reszta = nazwa.split("_LA_TH_", 1)
        rok = reszta[:4]
        czesci_po_roku = [p.strip() for p in reszta[4:].split(" - ") if p.strip()]
        typ = czesci_po_roku[0] if czesci_po_roku else ""
    else:
        prefiks = nazwa.split("_")[0]
        rok = str(datetime.datetime.now().year)
        typ = ""
    numer_fabryczny = nazwa.rsplit(" - ", 1)[-1]
    return prefiks, rok, numer_fabryczny, typ


# =============================================================================
# FORMATOWANIE DAT (POLSKIE) I OBSŁUGA DOKUMENTÓW WORD
# =============================================================================

def _formatuj_date(d):
    """Formatuje datetime/date jako 'DD miesiaca YYYY r.' (np. '04 maja 2026 r.')"""
    if isinstance(d, datetime.datetime):
        d = d.date()
    return f"{d.day:02d} {MIESIACE_GEN[d.month]} {d.year} r."


def _parsuj_date_dowolna(d):
    """
    Zwraca datetime.date z datetime/date albo ze STRINGA (np. '02.07.2026').
    None gdy nie da sie rozpoznac. Arkusze zapisuja date pomiaru w Strona 3 (kol. E)
    jako TEKST 'DD.MM.YYYY', dlatego musimy obsluzyc rowniez stringi.
    """
    if isinstance(d, datetime.datetime):
        return d.date()
    if isinstance(d, datetime.date):
        return d
    if isinstance(d, str):
        s = d.strip()
        if not s:
            return None
        for fmt in ('%d.%m.%Y', '%Y-%m-%d', '%d-%m-%Y', '%d.%m.%y', '%d/%m/%Y'):
            try:
                return datetime.datetime.strptime(s, fmt).date()
            except ValueError:
                continue
    return None


def _formatuj_daty_wzorcowania(daty_raw):
    """
    Formatuje liste dat jako polska date wzorcowania.
    Regula: ciagle 3+ dni → zakres (kreseczka); 1-2 dni → lista.
    Przy przejsciu miesiaca: nazwy obu miesiecy, rok tylko na koncu.

    Przyklady:
      [19.05, 20.05]         → '19, 20 maja 2026 r.'
      [10.05, 11.05, 12.05]  → '10 ÷ 12 maja 2026 r.'
      [29.04, 04.05]         → '29 kwietnia, 04 maja 2026 r.'
      [18.05, 19.05, 21.05]  → '18, 19, 21 maja 2026 r.'
    """
    daty = []
    for d in daty_raw:
        d = _parsuj_date_dowolna(d)
        if d is not None:
            daty.append(d)
    daty = sorted(set(daty))
    if not daty:
        return ""

    rok = daty[-1].year

    # Wyznacz segmenty: ciagle biegi >= 3 dni → zakres; reszta → pojedyncze daty
    segmenty = []
    i = 0
    while i < len(daty):
        j = i
        while j + 1 < len(daty) and (daty[j + 1] - daty[j]).days == 1:
            j += 1
        bieg = daty[i:j + 1]
        if len(bieg) >= 3:
            segmenty.append(("zakres", bieg))
        else:
            for d in bieg:
                segmenty.append(("pojedyncza", [d]))
        i = j + 1

    # Buduj pary (rok, miesiac, fragment_tekstowy)
    pary = []
    for typ_seg, daty_seg in segmenty:
        d0 = daty_seg[0]
        dN = daty_seg[-1]
        if typ_seg == "zakres" and d0.month == dN.month:
            pary.append((d0.year, d0.month, f"{d0.day:02d} ÷ {dN.day:02d}"))
        else:
            for d in daty_seg:
                pary.append((d.year, d.month, f"{d.day:02d}"))

    # Grupuj po (rok, miesiac) i scal fragmenty
    grupy = []
    for (yr, mo), grp in groupby(pary, key=lambda x: (x[0], x[1])):
        fragmenty = [item[2] for item in grp]
        grupy.append((yr, mo, ", ".join(fragmenty)))

    # Formatuj wynik: miesiac i rok tylko na ostatniej grupie
    czesci = []
    for idx, (yr, mo, dni_str) in enumerate(grupy):
        if idx < len(grupy) - 1:
            czesci.append(f"{dni_str} {MIESIACE_GEN[mo]}")
        else:
            czesci.append(f"{dni_str} {MIESIACE_GEN[mo]} {yr} r.")
    return ", ".join(czesci)


def _zastap_tekst_w_dok(doc, placeholder, value):
    """
    Zamienia placeholder we wszystkich paragrafach, tabelach i naglowkach/stopkach
    dokumentu Word. Obsluguje placeholder w pojedynczym runie jak i rozbity miedzy runami
    (skleja caly paragraf w pierwszym runie i podmienia tekst).
    """
    value = str(value) if value is not None else ""

    def _zastap_w_para(para):
        if placeholder not in para.text:
            return
        # Proba zamiany w pojedynczym runie (zachowuje formatowanie)
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return
        # Placeholder rozbity miedzy runami: znajdz minimalny zakres runow.
        # Uzywa mapowania pozycja-znaku -> indeks-runu, zeby nie niszczyc
        # elementow bez tekstu (np. <w:sym> dla symboli specjalnych).
        runs = para.runs
        texts = [r.text for r in runs]
        pelny = "".join(texts)
        pos = pelny.find(placeholder)
        if pos < 0:
            return
        pos_to_run = []
        for idx, t in enumerate(texts):
            pos_to_run.extend([idx] * len(t))
        if not pos_to_run or pos >= len(pos_to_run):
            return
        end_pos = pos + len(placeholder) - 1
        if end_pos >= len(pos_to_run):
            return
        i = pos_to_run[pos]
        j = pos_to_run[end_pos] + 1
        combined = "".join(texts[i:j])
        runs[i].text = combined.replace(placeholder, value)
        for r in runs[i + 1:j]:
            r.text = ""

    def _zastap_w_kontenerze(kontener):
        for para in kontener.paragraphs:
            _zastap_w_para(para)
        for table in kontener.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _zastap_w_para(para)

    _zastap_w_kontenerze(doc)

    # Naglowki i stopki wszystkich sekcji dokumentu
    for section in doc.sections:
        for hdr_ftr in (
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ):
            _zastap_w_kontenerze(hdr_ftr)


def _zastap_tekst_w_tabeli(table, placeholder, value):
    """
    Jak '_zastap_tekst_w_dok', ale ograniczone do jednej, konkretnej tabeli.
    Wymagane przez szablon 'zakres + temp.docx', ktory ma dwie niezalezne
    tabele kalibracyjne uzywajace tej samej numeracji placeholderow — podmiana
    w calym dokumencie nadpisalaby obie tabele tymi samymi wartosciami.
    """
    value = str(value) if value is not None else ""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if placeholder not in para.text:
                    continue
                zmieniono = False
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        zmieniono = True
                        break
                if zmieniono:
                    continue
                pelny = "".join(r.text for r in para.runs)
                if placeholder in pelny and para.runs:
                    para.runs[0].text = pelny.replace(placeholder, value)
                    for r in para.runs[1:]:
                        r.text = ""


def _zastap_tekst_w_tabeli_z_sup_po(table, placeholder, value, sup_text):
    """Like _zastap_tekst_w_tabeli but appends sup_text as superscript run after replaced value."""
    value_str = str(value) if value is not None else ""

    def _dodaj_run_sup(after_run):
        new_r = OxmlElement('w:r')
        src_rPr = after_run._r.find(qn('w:rPr'))
        new_rPr = deepcopy(src_rPr) if src_rPr is not None else OxmlElement('w:rPr')
        for va in new_rPr.findall(qn('w:vertAlign')):
            new_rPr.remove(va)
        va_el = OxmlElement('w:vertAlign')
        va_el.set(qn('w:val'), 'superscript')
        new_rPr.append(va_el)
        new_r.append(new_rPr)
        wt = OxmlElement('w:t')
        wt.text = sup_text
        if sup_text and sup_text[0] == ' ':
            wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        new_r.append(wt)
        after_run._r.addnext(new_r)

    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if placeholder not in para.text:
                    continue
                zmieniono = False
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value_str)
                        _dodaj_run_sup(run)
                        zmieniono = True
                        break
                if zmieniono:
                    continue
                pelny = "".join(r.text for r in para.runs)
                if placeholder in pelny and para.runs:
                    para.runs[0].text = pelny.replace(placeholder, value_str)
                    for r in para.runs[1:]:
                        r.text = ""
                    _dodaj_run_sup(para.runs[0])


def _zastap_histereza_w_dok(doc, punkty):
    """Replaces [histereza] with superscript '1)' + note text about hysteresis measurement.
    Finds the hysteresis point (nazwa ending in ' (N)') in punkty, rounds T to integer
    and RH to nearest 10, then builds: ^1) Punkt (T °C, H %rh) powtórzony...
    If no hysteresis point, removes the placeholder."""
    PLACEHOLDER = "[histereza]"
    hist = next((p for p in punkty if _to_jest_powtorka_kolizji(p.get("nazwa"))), None)

    if hist is None:
        _zastap_tekst_w_dok(doc, PLACEHOLDER, "")
        return

    try:
        temp_r = int(_round_half_away_from_zero(float(hist.get("wartosc_odn"))))
    except (TypeError, ValueError):
        temp_r = "?"
    try:
        rh_r = int(_round_half_away_from_zero(float(hist.get("wartosc_odn_RH")) / 10)) * 10
    except (TypeError, ValueError):
        rh_r = "?"

    tekst_glowny = (
        f" Punkt ({temp_r} °C, {rh_r} %rh) powtórzony"
        f" w celu wyznaczenia wartości histerezy wskazań"
        f" wzorcowanego przyrządu."
    )

    def _podmien_w_para(para):
        if PLACEHOLDER not in para.text:
            return
        for run in para.runs:
            if PLACEHOLDER not in run.text:
                continue
            before = run.text[:run.text.index(PLACEHOLDER)]
            after = run.text[run.text.index(PLACEHOLDER) + len(PLACEHOLDER):]
            run.text = before + "1)"
            run.font.superscript = True
            main_r = OxmlElement('w:r')
            src_rPr = run._r.find(qn('w:rPr'))
            main_rPr = deepcopy(src_rPr) if src_rPr is not None else OxmlElement('w:rPr')
            for va in main_rPr.findall(qn('w:vertAlign')):
                main_rPr.remove(va)
            main_r.append(main_rPr)
            wt = OxmlElement('w:t')
            wt.text = tekst_glowny + after
            wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            main_r.append(wt)
            run._r.addnext(main_r)
            return
        pelny = "".join(r.text for r in para.runs)
        if PLACEHOLDER not in pelny or not para.runs:
            return
        pos = pelny.index(PLACEHOLDER)
        before = pelny[:pos]
        after = pelny[pos + len(PLACEHOLDER):]
        for r in para.runs:
            r.text = ""
        r0 = para.runs[0]
        r0.text = before + "1)"
        r0.font.superscript = True
        main_r = OxmlElement('w:r')
        src_rPr = r0._r.find(qn('w:rPr'))
        main_rPr = deepcopy(src_rPr) if src_rPr is not None else OxmlElement('w:rPr')
        for va in main_rPr.findall(qn('w:vertAlign')):
            main_rPr.remove(va)
        main_r.append(main_rPr)
        wt = OxmlElement('w:t')
        wt.text = tekst_glowny + after
        wt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        main_r.append(wt)
        r0._r.addnext(main_r)

    def _zastap_w_kontenerze(kontener):
        for para in kontener.paragraphs:
            _podmien_w_para(para)
        for table in kontener.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _podmien_w_para(para)

    _zastap_w_kontenerze(doc)
    for section in doc.sections:
        for hdr_ftr in (
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ):
            _zastap_w_kontenerze(hdr_ftr)


_WZORZEC_PLACEHOLDER_PIERWSZY = re.compile(r'\[(?:wartość_odn|wartosc_odn)_1\]')
_WZORZEC_PLACEHOLDER_PIERWSZY_RH = re.compile(r'\[(?:wartość_odn|wartosc_odn)_1_RH\]')


def _znajdz_tabele_kalibracji(doc):
    """
    Zwraca liste tabel kalibracyjnych znalezionych w dokumencie (po obecnosci
    placeholdera '[wartość_odn_1]' w jednej z komorek). Szablony 'zakres.docx'
    i 'tylko temp.docx' maja jedna taka tabele; szablon 'zakres + temp.docx'
    (przypadek mieszany — czesc punktow z wilgotnoscia, czesc bez) ma dwie,
    niezalezne, kazda z wlasna numeracja placeholderow zaczynajaca sie od 1.
    """
    tabele = []
    for table in doc.tables:
        znaleziono = False
        for row in table.rows:
            for cell in row.cells:
                if _WZORZEC_PLACEHOLDER_PIERWSZY.search(cell.text):
                    znaleziono = True
                    break
            if znaleziono:
                break
        if znaleziono:
            tabele.append(table)
    return tabele


def _punkty_z_rh(punkty):
    """Tylko punkty odczytane z wilgotnoscia (osiem pol, w tym 'wartosc_odn_RH')."""
    return [p for p in punkty if "wartosc_odn_RH" in p]


def _punkty_temp_only_zdeduplikowane(punkty):
    """
    Wszystkie punkty zredukowane do jednego reprezentanta per zaokraglona
    temperatura — analogicznie do tabeli 'tylko temperatura' (O1:R9) w arkuszu
    Wyniki: gdy ten sam nominal wystepuje dla kilku zakladek (np. rozne celowe
    wilgotnosci przy tej samej temperaturze), wybierany jest jeden punkt:
    bez sufiksu ' (N)' (czyli nie powtorka histerezy), a z tych — najblizszy
    50% RH (lub pierwszy, gdy zaden punkt w grupie nie ma wilgotnosci).
    """
    wpisy = []
    for p in punkty:
        try:
            temp = _round_half_away_from_zero(float(p.get("wartosc_odn")))
        except (TypeError, ValueError):
            continue
        # Nominal RH z nazwy zakladki (pewniejszy niz odczyt) — do wyboru ~50%.
        rh = _rh_z_nazwy_zakladki(p.get("nazwa"))
        powtorka = _to_jest_powtorka_kolizji(p.get("nazwa"))
        wpisy.append((temp, rh, powtorka, p))
    return _wybierz_reprezentantow_temp(wpisy)


def _uzupelnij_tabele_kalibracji(doc, punkty):
    """
    Wyszukuje tabele kalibracyjna/e w dokumencie (po '[wartość_odn_1]') i wypelnia
    kazda znaleziona tabele odpowiednim podzbiorem 'punkty':
      - tabela z placeholderami '[..._1_RH]'  -> tylko punkty z wilgotnoscia
        ('_punkty_z_rh'),
      - tabela bez nich (tylko temperatura)   -> wszystkie punkty, zredukowane
        do jednego reprezentanta per temperatura ('_punkty_temp_only_zdeduplikowane').
    Szablony 'zakres.docx' / 'tylko temp.docx' maja jedna tabele (zachowanie
    jak dawniej — dla 'zakres.docx' wszystkie punkty mialy wilgotnosc, dla
    'tylko temp.docx' zaden punkt jej nie mial). Szablon 'zakres + temp.docx'
    (przypadek mieszany) ma dwie niezalezne tabele wypelniane oddzielnie.

    punkty = [{"nazwa": ..., "wartosc_odn": ..., "zmierzona": ..., "poprawka": ...,
               "niepewnosc": ..., "wartosc_odn_RH": ..., "zmierzona_RH": ...,
               "poprawka_RH": ..., "niepewnosc_RH": ...}, ...]
    Pola "*_RH" sa obecne tylko dla punktow odczytanych z zakladki majacej
    aktywna wilgotnosc.
    """
    tabele = _znajdz_tabele_kalibracji(doc)
    if not tabele:
        return  # brak tabeli kalibracyjnej w dokumencie

    for table in tabele:
        ma_rh_tabeli = any(
            _WZORZEC_PLACEHOLDER_PIERWSZY_RH.search(cell.text)
            for row in table.rows for cell in row.cells
        )
        punkty_tabeli = _punkty_z_rh(punkty) if ma_rh_tabeli else _punkty_temp_only_zdeduplikowane(punkty)
        _wypelnij_jedna_tabele_kalibracji(table, punkty_tabeli)


def _wypelnij_jedna_tabele_kalibracji(cal_table, punkty):
    """
    Dopasowuje liczbe wierszy placeholderow danej tabeli do liczby punktow
    pomiarowych (dodaje klonujac ostatni istniejacy wiersz placeholderow,
    usuwa nadmiarowe), a nastepnie wypelnia wszystkie jej komorki wartosciami
    z 'punkty'. Podmiana tekstu jest ograniczona do tej jednej tabeli (patrz
    '_zastap_tekst_w_tabeli') — przy szablonie z dwiema tabelami obie uzywaja
    tej samej numeracji placeholderow.

    Punkty sa sortowane rosnaco po wartosc_odn (zaokraglonej do stopnia) przed
    wypelnieniem; gdy dwa punkty zaokraglaja sie do tego samego stopnia, "wyzszy"
    jest ten z mniejsza wilgotnoscia (wartosc_odn_RH) — umieszczany nizej w tabeli.

    Wiersze placeholderow sa wykrywane dynamicznie (wg numeru w '[wartość_odn_N]'),
    a nie po stalym indeksie wiersza — szablon moze miec wiecej niz 3 wbudowane
    wiersze (np. po recznym powiekszeniu przez uzytkownika), a stary kod ze stalym
    indeksem 2 trafial w wiersz jednostek ('°C'), nie w trzeci wiersz danych.

    Pogrubiona ramka tabeli w szablonie jest ustawiona jako formatowanie
    bezposrednie (tcBorders/bottom) na komorkach fizycznie ostatniego wiersza,
    nie na poziomie calej tabeli:
      - przy USUWANIU wierszy jest ona przenoszona na nowy ostatni wiersz,
      - przy DODAWANIU wierszy (klonowanie) kazdy klon dziedziczy te ramke
        po zrodle — bez korekty zostalaby ona blednie na KAZDYM sklonowanym
        wierszu (w tym na oryginalnym, ktory przestal byc ostatni), wiec po
        dodaniu wszystkich wierszy jest usuwana ze wszystkich oprocz
        faktycznie ostatniego.
    """
    def _klucz_sortowania(p):
        v = p.get("wartosc_odn")
        try:
            temp = float(v)
        except (TypeError, ValueError):
            return (float('inf'), 1, float('inf'), float('inf'))
        rh = p.get("wartosc_odn_RH")
        try:
            rh_f = float(rh)
        except (TypeError, ValueError):
            rh_f = 0.0
        # Powtorki histerezy (sufiks ' (N)') na koncu grupy temperaturowej;
        # w ramach tej samej grupy RH rosnie (10 → 50 → 85 → 50(2)).
        is_dup = 1 if _to_jest_powtorka_kolizji(p.get("nazwa")) else 0
        return (_round_half_away_from_zero(temp), is_dup, rh_f, temp)

    punkty = sorted(punkty, key=_klucz_sortowania)

    n = len(punkty)

    # Znajdz wiersze placeholderow wg numeru w "[wartość_odn_<k>]" / "[wartosc_odn_<k>]"
    wzorzec_nr = re.compile(r'\[(?:wartość_odn|wartosc_odn)_(\d+)\]')
    wiersz_po_numerze = {}
    for ri, row in enumerate(cal_table.rows):
        for cell in row.cells:
            m = wzorzec_nr.search(cell.text)
            if m:
                wiersz_po_numerze[int(m.group(1))] = ri
                break
    if not wiersz_po_numerze:
        return  # brak wykrywalnych wierszy placeholderow

    numery = sorted(wiersz_po_numerze.keys())
    m_istniejace = len(numery)

    def _pobierz_dolne_obramowanie(row):
        """Zwraca per-komorkowa kopie elementu <w:bottom> z tcBorders danego wiersza
        (lub None, gdy komorka nie ma wlasnego bezposredniego obramowania dolnego)."""
        wynik = []
        for cell in row.cells:
            dolna = None
            tcPr = cell._tc.tcPr
            if tcPr is not None:
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    el = tcBorders.find(qn('w:bottom'))
                    if el is not None:
                        dolna = deepcopy(el)
            wynik.append(dolna)
        return wynik

    def _ustaw_dolne_obramowanie(row, dolne_komorek):
        for cell, dolna in zip(row.cells, dolne_komorek):
            if dolna is None:
                continue
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.insert_element_before(
                    tcBorders, 'w:shd', 'w:noWrap', 'w:tcMar', 'w:textDirection',
                    'w:tcFitText', 'w:vAlign', 'w:hideMark', 'w:headers',
                    'w:cellIns', 'w:cellDel', 'w:cellMerge', 'w:tcPrChange')
            istniejaca = tcBorders.find(qn('w:bottom'))
            if istniejaca is not None:
                tcBorders.remove(istniejaca)
            tcBorders.append(deepcopy(dolna))

    def _wyczysc_dolne_obramowanie(row):
        for cell in row.cells:
            tcPr = cell._tc.tcPr
            if tcPr is None:
                continue
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                continue
            istniejaca = tcBorders.find(qn('w:bottom'))
            if istniejaca is not None:
                tcBorders.remove(istniejaca)

    # Dodaj brakujace wiersze (klonujac ostatni istniejacy wiersz placeholderow)
    if n > m_istniejace:
        wzorzec_przenumer_bazowy = lambda nr: re.compile(rf'_{nr}(_RH)?\]')
        ostatni_nr = numery[-1]
        ostatni_idx = wiersz_po_numerze[ostatni_nr]
        dolne_obramowanie_oryg = _pobierz_dolne_obramowanie(cal_table.rows[ostatni_idx])
        wiersze_do_oczyszczenia = []
        for k in range(m_istniejace, n):
            wiersze_do_oczyszczenia.append(ostatni_idx)
            src_tr = cal_table.rows[ostatni_idx]._tr
            new_tr = deepcopy(src_tr)
            cal_table._tbl.append(new_tr)
            new_row = cal_table.rows[-1]
            wzorzec_przenumer = wzorzec_przenumer_bazowy(ostatni_nr)
            for cell in new_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = wzorzec_przenumer.sub(
                            lambda m: f"_{k + 1}{m.group(1) or ''}]", run.text)
            ostatni_idx = len(cal_table.rows) - 1
            ostatni_nr = k + 1
        # Ramka dolna sklonowala sie na kazdy nowy wiersz (i pozostala na
        # zrodlowym, ktory przestal byc ostatni) — usun ja wszedzie oprocz
        # faktycznie ostatniego wiersza, na ktorym przywracamy oryginalna wartosc.
        for idx in wiersze_do_oczyszczenia:
            _wyczysc_dolne_obramowanie(cal_table.rows[idx])
        _ustaw_dolne_obramowanie(cal_table.rows[-1], dolne_obramowanie_oryg)

    # Usun nadmiarowe wiersze placeholderow (gdy punktow jest mniej niz w szablonie).
    # Usuwamy od najwyzszego indeksu do najnizszego — usuniecie wiersza przesuwa
    # indeksy wszystkich pozniejszych wierszy o 1, wiec idac od konca unikamy
    # uzycia juz nieaktualnego (przestarzalego) indeksu.
    if n < m_istniejace:
        dolne_obramowanie_oryg = _pobierz_dolne_obramowanie(cal_table.rows[-1])
        indeksy_do_usuniecia = sorted((wiersz_po_numerze[nr] for nr in numery[n:]), reverse=True)
        for idx in indeksy_do_usuniecia:
            tr = cal_table.rows[idx]._tr
            tr.getparent().remove(tr)
        _ustaw_dolne_obramowanie(cal_table.rows[-1], dolne_obramowanie_oryg)

    def _fmt(v):
        if v is None:
            return ""
        if isinstance(v, float) and v == int(v):
            return str(int(v))
        return str(v).replace(".", ",")  # separator dziesiętny: przecinek (PL)

    for k, punkt in enumerate(punkty):
        idx = k + 1
        val_odn = _fmt(punkt.get("wartosc_odn"))
        val_zmi = _fmt(punkt.get("zmierzona"))
        val_pop = _fmt(punkt.get("poprawka"))
        val_nie = _fmt(punkt.get("niepewnosc"))
        val_odn_rh = _fmt(punkt.get("wartosc_odn_RH"))
        val_zmi_rh = _fmt(punkt.get("zmierzona_RH"))
        val_pop_rh = _fmt(punkt.get("poprawka_RH"))
        val_nie_rh = _fmt(punkt.get("niepewnosc_RH"))
        is_hist = _to_jest_powtorka_kolizji(punkt.get("nazwa"))
        # Probujemy obie wersje: z polskimi znakami i bez (ASCII), dla pewnosci
        for placeholder, wartosc in [
            (f"[wartość_odn_{idx}]", val_odn),
            (f"[wartosc_odn_{idx}]",  val_odn),
            (f"[zmierzona_{idx}]",    val_zmi),
            (f"[poprawka_{idx}]",     val_pop),
            (f"[niepewność_{idx}]",  val_nie),
            (f"[niepewnosc_{idx}]",   val_nie),
            (f"[wartość_odn_{idx}_RH]", val_odn_rh),
            (f"[wartosc_odn_{idx}_RH]",  val_odn_rh),
            (f"[zmierzona_{idx}_RH]",    val_zmi_rh),
            (f"[poprawka_{idx}_RH]",     val_pop_rh),
            (f"[niepewność_{idx}_RH]",  val_nie_rh),
            (f"[niepewnosc_{idx}_RH]",   val_nie_rh),
        ]:
            # Dla punktu histerezy: wartość odn. RH dostaje 3 spacje + górny indeks " 1)"
            if is_hist and placeholder in (
                f"[wartość_odn_{idx}_RH]", f"[wartosc_odn_{idx}_RH]"
            ):
                _zastap_tekst_w_tabeli_z_sup_po(cal_table, placeholder, "   " + wartosc, " 1)")
            else:
                _zastap_tekst_w_tabeli(cal_table, placeholder, wartosc)


def generuj_nazwe_word(nr_sw, prefiks, rok):
    """Generuje nazwe pliku Word: '{nr_sw}_{prefiks}_LA_TH_{rok}.docx'"""
    return f"{nr_sw}_{prefiks}_LA_TH_{rok}.docx"


# PZ jako fallback danych przyrzadu (gdy Strona 2 pusta — np. protokol bez PZ-fill).
PZ_FOLDER_ARK = C.sciezka("CC_PZ_FOLDER", "PZ", FOLDER)
_PZ_CACHE = None

def _pz_mapa_arkusze():
    """Leniwie wczytuje mape PZ (po nr fabrycznym); cache na czas dzialania."""
    global _PZ_CACHE
    if _PZ_CACHE is None:
        _PZ_CACHE, _ = pz_dane.wczytaj_pz(PZ_FOLDER_ARK)
    return _PZ_CACHE


_PZ_PUNKTY_CACHE = None


def _pz_punkty_arkusze():
    """
    Leniwie wczytuje punkty ZAMOWIONE w PZ (wartosci nominalne, np. (25, 30));
    cache na czas dzialania. Sluza do nazywania zakladek arkusza obliczeniowego.
    """
    global _PZ_PUNKTY_CACHE
    if _PZ_PUNKTY_CACHE is None:
        _PZ_PUNKTY_CACHE = pz_dane.wczytaj_punkty(PZ_FOLDER_ARK)
    return _PZ_PUNKTY_CACHE


# Tolerancje dopasowania nastawy komory do punktu NOMINALNEGO z PZ
# (np. PZ '25 °C / 30 %' -> nastawa komory 25,0 / 28,0).
TOL_NAZWA_T  = 1.5    # [st.C]
TOL_NAZWA_RH = 4.0    # [%RH]


def _nazwy_zakladek_z_pz(dane_zakladek):
    """
    Podmienia nazwy zakladek na wartosci NOMINALNE z PZ ('25, 30' zamiast '25, 28').

    Nastawa komory bywa inna niz punkt zamowiony (25/28 zamiast 25/30, 24,9/48,5 zamiast
    25/50), a w dokumentach obowiazuje punkt z zamowienia.

    Dopasowanie jest PER BLOK — do kazdego bloku szukamy NAJBLIZSZEGO jeszcze nieuzytego
    punktu z PZ. Liczby nie musza sie zgadzac: PZ obejmuje czesto wiecej punktow, niz
    odbyto w danym wsadzie komory (np. 13 zamowionych, 4 wykonane), a przyrzady z jednego
    PZ bywaja wzorcowane w kilku wsadach. Blok bez pasujacego punktu zachowuje nazwe
    z nastawy komory.
    """
    punkty = _pz_punkty_arkusze()
    if not punkty or not dane_zakladek:
        return

    uzyte, zmienione = set(), 0
    for zd in dane_zakladek:
        t_n, rh_n = zd.get("nast_t"), zd.get("nast_rh")
        if t_n is None:
            continue
        najlepszy, najlepszy_dyst = None, None
        for idx, (t_pz, rh_pz) in enumerate(punkty):
            if idx in uzyte or t_pz is None:
                continue
            if abs(t_n - t_pz) > TOL_NAZWA_T:
                continue
            # punkt tylko-temperaturowy w PZ vs nastawa RH=0/'-' w protokole
            if rh_pz is None:
                if rh_n not in (None, 0.0):
                    continue
                dyst = abs(t_n - t_pz)
            else:
                if rh_n is None or abs(rh_n - rh_pz) > TOL_NAZWA_RH:
                    continue
                dyst = abs(t_n - t_pz) + abs(rh_n - rh_pz) / 10.0
            if najlepszy_dyst is None or dyst < najlepszy_dyst - 1e-9:
                najlepszy, najlepszy_dyst = idx, dyst
        if najlepszy is None:
            continue
        uzyte.add(najlepszy)
        t_pz, rh_pz = punkty[najlepszy]
        czesc_t  = _cell_to_sheet_name_part(t_pz)
        czesc_rh = _cell_to_sheet_name_part(rh_pz) if rh_pz is not None else "-"
        zd["nazwa"] = f"{czesc_t}, {czesc_rh}"
        zmienione += 1

    if zmienione:
        print(f"  Nazwy zakladek wg punktow z PZ ({zmienione}/{len(dane_zakladek)}): "
              + " | ".join(f"'{zd['nazwa']}'" for zd in dane_zakladek))
    else:
        _warn("Zadnego bloku protokolu nie udalo sie dopasowac do punktow z PZ — "
              "nazwy zakladek zostaja z nastaw komory.")


def _wariant_uzytkownik(nazwa_szablonu):
    """'... tylko temp.docx' -> '... tylko temp (uzytkownik).docx' (wariant z UZYTKOWNIKIEM)."""
    baza, ext = os.path.splitext(nazwa_szablonu)
    return f"{baza} (uzytkownik){ext}"


def utworz_kopie_word(folder, szablon_word_tylko_temp, szablon_word_z_rh, szablon_word_mieszany, dane_s2, kopie_excel,
                      dane_zakladek, dane_kalibracji, nr_sw_poczatkowy,
                      dane_zakladek_per_kopia=None, klasa_wilg_per_kopia=None,
                      warunki_per_kopia=None):
    """
    Dla kazdej kopii Excel tworzy kopie dokumentu Word i wypelnia placeholdery:
      [data]             aktualna data ('DD miesiaca YYYY r.')
      [nr_sw]            numer swiadectwa (nr_sw_poczatkowy + j)
      [nr_zl]            nr zamowienia (z rekordu Strona 2 kol. D → E6 kopii)
      [nr_fabr]          nr fabryczny  (z nazwy pliku → G6 kopii)
      [nr_ewid]          nr ewidencyjny (z rekordu Strona 2 kol. F)
      [wytworca]         wytworca      (z rekordu Strona 2 kol. B → E5 kopii)
      [typ]              typ/model     (segment po roku w nazwie pliku)
      [data_wzorcowania] daty pomiaru  (K4_val zakładek, sformatowane)
      [Podpis]           PODPISUJACY_2
      tabela/e kalibracyjna/e: per zakladka robocza, z D246:G246 (punkty bez
        wilgotnosci) i/lub D239:K239 (punkty z wilgotnoscia) — patrz
        '_uzupelnij_tabele_kalibracji'.
    Szablon Word jest wybierany per kopia wg klasa_wilg_per_kopia[j]:
      "brak"     -> szablon_word_tylko_temp (zadna zakladka nie ma wilgotnosci)
      "pelna"    -> szablon_word_z_rh       (wszystkie zakladki maja wilgotnosc)
      "mieszana" -> szablon_word_mieszany   (czesc zakladek ma wilgotnosc, czesc nie —
                    szablon z dwiema niezaleznymi tabelami kalibracyjnymi)
    Zwraca liste nazw utworzonych plikow Word.
    """
    if not _DOCX_OK:
        print("  [BLAD] Brak biblioteki 'python-docx'. Zainstaluj: pip install python-docx")
        return []

    n = len(kopie_excel)
    nazwy_word = []

    dzis = datetime.datetime.now()
    data_str = _formatuj_date(dzis)

    for j, (rekord, nowa_nazwa_xlsx) in enumerate(zip(dane_s2, kopie_excel)):
        nr_sw = nr_sw_poczatkowy + j
        prefiks, rok, nr_fab, typ = _parsuj_nazwe_pliku(nowa_nazwa_xlsx)

        # Dane przyrzadu z PZ (po nr fabrycznym): fallback pol Strony 2 + ew. UZYTKOWNIK.
        _pzdev = _pz_mapa_arkusze().get(pz_dane.normalizuj_serial(nr_fab))
        uzytkownik_v = (_pzdev.uzytkownik if _pzdev else "") or ""
        # Blok adresowy ZLECENIODAWCY z PZ (nazwa + ulica + kod/miasto), wieloliniowy.
        zleceniodawca_v = (getattr(_pzdev, "zleceniodawca", "") if _pzdev else "") or ""
        wytworca_v = str(rekord.get("B") or "") or (_pzdev.wytworca if _pzdev else "")
        typ_v      = str(rekord.get("D") or "") or (_pzdev.typ if _pzdev else "")
        nr_ewid_v  = _cell_to_str(rekord.get("F")) or (_pzdev.nr_ewid if _pzdev else "")

        klasa = klasa_wilg_per_kopia[j] if klasa_wilg_per_kopia and j < len(klasa_wilg_per_kopia) else "brak"
        if klasa == "mieszana":
            szablon_word = szablon_word_mieszany or szablon_word_z_rh or szablon_word_tylko_temp
        elif klasa == "pelna":
            szablon_word = szablon_word_z_rh or szablon_word_tylko_temp
        else:
            szablon_word = szablon_word_tylko_temp or szablon_word_z_rh
        if not szablon_word:
            print(f"  [UWAGA] Brak skonfigurowanego szablonu Word dla kopii '{nowa_nazwa_xlsx}' — pomijam.")
            continue

        # Gdy PZ zawiera UZYTKOWNIKA — uzyj wariantu szablonu '(uzytkownik)', jesli istnieje.
        if uzytkownik_v:
            wariant = _wariant_uzytkownik(szablon_word)
            if os.path.exists(os.path.join(folder, wariant)):
                szablon_word = wariant
                print(f"  [Word] Kopia '{nowa_nazwa_xlsx}': UZYTKOWNIK obecny -> szablon '{wariant}'.")
            else:
                print(f"  [UWAGA] Brak szablonu '{wariant}' — uzywam standardowego (bez uzytkownika).")

        szablon_path = os.path.join(folder, szablon_word)
        if not os.path.exists(szablon_path):
            print(f"  [UWAGA] Plik szablonu Word nie istnieje: {szablon_word} — pomijam kopie '{nowa_nazwa_xlsx}'.")
            continue

        nowa_nazwa_docx = generuj_nazwe_word(nr_sw, prefiks, rok)
        sciezka_docx = os.path.join(folder, nowa_nazwa_docx)

        if dane_zakladek_per_kopia is not None and j < len(dane_zakladek_per_kopia):
            zakladki_do_daty = dane_zakladek_per_kopia[j]
        else:
            zakladki_do_daty = dane_zakladek
        daty_k4 = [zd.get("K4_val") for zd in zakladki_do_daty]
        data_wzorcowania = _formatuj_daty_wzorcowania(daty_k4)

        shutil.copy2(szablon_path, sciezka_docx)
        doc = DocxDocument(sciezka_docx)

        for placeholder, wartosc in {
            "[data]":             data_str,
            "[nr_sw]":            str(nr_sw),
            "[nr_zl]":            prefiks,          # prefiks z nazwy pliku (przed _LA_TH_)
            "[nr_fabr]":          nr_fab,
            "[nr_ewid]":          nr_ewid_v,        # Strona 2 kol. F (lub PZ)
            "[wytworca]":         wytworca_v,       # Strona 2 kol. B (lub PZ)
            "[typ]":              typ_v,            # Strona 2 kol. D (lub PZ)
            "[data_wzorcowania]": data_wzorcowania,
            "[użytkownik]":       uzytkownik_v,     # blok adresowy z PZ (tylko szablon uzytkownika)
            "[zleceniodawca]":    zleceniodawca_v,  # blok adresowy ZLECENIODAWCY z PZ
            "[Podpis]":           PODPISUJACY_2,
        }.items():
            _zastap_tekst_w_dok(doc, placeholder, wartosc)

        punkty = dane_kalibracji[j] if j < len(dane_kalibracji) else []
        _uzupelnij_tabele_kalibracji(doc, punkty)
        _zastap_histereza_w_dok(doc, punkty)

        zakresy = warunki_per_kopia[j] if warunki_per_kopia and j < len(warunki_per_kopia) else None
        if zakresy is None:
            zakresy = {}
        # Temperatura otoczenia z jednym miejscem po przecinku, wilgotnosc —
        # w pelnych procentach (bez ',0').
        for placeholder, klucz, miejsca in [
            ("[temp_min]", "temp_min", 1),
            ("[temp_max]", "temp_max", 1),
            ("[wilg_min]", "wilg_min", 0),
            ("[wilg_max]", "wilg_max", 0),
        ]:
            _zastap_tekst_w_dok(doc, placeholder,
                                _formatuj_zakres_srodowiskowy(zakresy.get(klucz), miejsca))

        doc.save(sciezka_docx)
        nazwy_word.append(nowa_nazwa_docx)
        print(f"  [Word {j + 1:>{len(str(n))}}/{n}] {nowa_nazwa_docx}")

    return nazwy_word


# =============================================================================
# ZARZĄDZANIE ZAKŁADKAMI I WYPEŁNIANIE KOMÓREK
# =============================================================================

def _dopasuj_liczbe_zakładek(wb, working, n, ark_wyniki, first_ws_name):
    """
    Usuwa nadmiarowe zakładki robocze (od końca) lub dodaje kopie
    pierwszej zakładki roboczej gdy brakuje.
    Zwraca zaktualizowaną listę nazw zakładek roboczych.
    """
    if len(working) > n:
        # Usuń nadmiarowe od końca listy
        for ws_name in working[n:]:
            del wb[ws_name]
        working = working[:n]

    elif len(working) < n:
        # Dodaj kopie pierwszej zakładki roboczej
        first_ws = wb[first_ws_name]
        for _ in range(n - len(working)):
            new_ws = wb.copy_worksheet(first_ws)
            # Przesuń nową zakładkę o jedną pozycję w lewo (tuż przed Wyniki)
            if ark_wyniki in wb.sheetnames:
                wb.move_sheet(new_ws.title, offset=-1)
            working.append(new_ws.title)

    return working


def dostosuj_zakladki_i_wypelnij(sciezka_pliku, dane_zakladek, dane_ef_kopia, ark_wyniki):
    """
    Otwiera skopiowany plik xlsx i:
      1. Dopasowuje liczbę zakładek roboczych do len(dane_zakladek).
      2. Zmienia nazwy zakładek roboczych zgodnie z dane_zakladek[i]["nazwa"].
      3. Wypełnia w każdej zakładce roboczej:
           C15:C19 ← dane_zakladek[i]["L_dane"]
           D15:D19 ← dane_zakladek[i]["M_dane"]
           E15:E19 ← dane_ef_kopia[i]["E_dane"]
           F15:F19 ← dane_ef_kopia[i]["F_dane"]
      4. Zapisuje plik (jedno wb.save() na końcu).
    Zakładka ARKUSZ_WYNIKI ("Wyniki") nie jest nigdy modyfikowana.
    """
    wb = openpyxl.load_workbook(sciezka_pliku, keep_links=False)

    wyniki_ok = ark_wyniki in wb.sheetnames
    if not wyniki_ok:
        print(f"  [OSTRZEŻENIE] Brak zakładki '{ark_wyniki}' w pliku.")

    # Pobierz zakładki robocze (wszystkie poza chronioną)
    working = [s for s in wb.sheetnames if s != ark_wyniki]

    if not working:
        print(f"  [BŁĄD] Brak zakładek roboczych — pomijam plik.")
        wb.close()
        return

    if not dane_zakladek:
        print(f"  [OSTRZEŻENIE] Brak definicji zakładek — pomijam modyfikację zakładek.")
        wb.close()
        return

    N = len(dane_zakladek)
    first_ws_name = working[0]

    # --- Etap 2: dopasuj liczbę zakładek roboczych ---
    working = _dopasuj_liczbe_zakładek(wb, working, N, ark_wyniki, first_ws_name)

    # Odśwież listę po zmianach struktury
    working = [s for s in wb.sheetnames if s != ark_wyniki]

    # --- Etap 2: zmień nazwy przez tymczasowe (unika konfliktów między starymi a nowymi) ---
    for i, ws_name in enumerate(working):
        wb[ws_name].title = f"__tmp_{i}__"

    docelowe_nazwy = _unikalne_nazwy_zakladek(dane_zakladek)
    working_tmp = [s for s in wb.sheetnames if s != ark_wyniki]
    for i, ws_name in enumerate(working_tmp):
        wb[ws_name].title = docelowe_nazwy[i]

    # --- Etap 3 & 4: wypełnij komórki w każdej zakładce roboczej ---
    working_final = [s for s in wb.sheetnames if s != ark_wyniki]

    ADRESY_C = ["C15", "C16", "C17", "C18", "C19"]
    ADRESY_D = ["D15", "D16", "D17", "D18", "D19"]
    ADRESY_E = ["E15", "E16", "E17", "E18", "E19"]
    ADRESY_F = ["F15", "F16", "F17", "F18", "F19"]

    for i, ws_name in enumerate(working_final):
        ws = wb[ws_name]
        zd = dane_zakladek[i]
        ef = dane_ef_kopia[i] if i < len(dane_ef_kopia) else None

        # C15:C19 ← domyslnie L_dane; dla CC-04 per-kopia C_dane.
        for offset, addr in enumerate(ADRESY_C):
            val = None
            if isinstance(ef, dict) and "C_dane" in ef:
                c_dane = ef.get("C_dane")
                if isinstance(c_dane, list) and offset < len(c_dane):
                    val = c_dane[offset]
            elif val is None:
                val = zd["L_dane"][offset]
            if val is not None:
                ws[addr] = val

        # D15:D19 ← domyslnie M_dane; dla CC-04 per-kopia D_dane.
        for offset, addr in enumerate(ADRESY_D):
            val = None
            if isinstance(ef, dict) and "D_dane" in ef:
                d_dane = ef.get("D_dane")
                if isinstance(d_dane, list) and offset < len(d_dane):
                    val = d_dane[offset]
            elif val is None:
                val = zd["M_dane"][offset]
            if val is not None:
                ws[addr] = val

        # E15:E19 i F15:F19 ← dane zależne od numeru kopii
        if isinstance(ef, dict):
            for offset, addr in enumerate(ADRESY_E):
                val = ef.get("E_dane", [None] * len(ADRESY_E))[offset]
                if val is not None:
                    ws[addr] = val
            for offset, addr in enumerate(ADRESY_F):
                val = ef.get("F_dane", [None] * len(ADRESY_F))[offset]
                if val is not None:
                    ws[addr] = val

    wb.save(sciezka_pliku)


# =============================================================================
# ZAPIS DO KOPII PRZEZ XLWINGS (COM Excel) — zachowuje formuly
# =============================================================================

_COLS_BI = ["B", "C", "D", "E", "F", "G", "H", "I"]
_COLS_OR = ["O", "P", "Q", "R"]


def _czytaj_wiersz_formul(ws_w, r, kolumny):
    return [ws_w.range(f"{c}{r}").formula for c in kolumny]


def _zapisz_wiersz_formul(ws_w, r, kolumny, formuly):
    for c, f in zip(kolumny, formuly):
        if f:
            ws_w.range(f"{c}{r}").formula = f
        else:
            ws_w.range(f"{c}{r}").clear_contents()


def _wyczysc_wiersz(ws_w, r, kolumny):
    ws_w.range(f"{kolumny[0]}{r}:{kolumny[-1]}{r}").clear_contents()


_WZORZEC_ARKUSZA_W_FORMULE = re.compile(r"^=(?:'([^']+)'|([^!'][^!]*))!")


def _arkusz_z_formuly(formula_text):
    if not isinstance(formula_text, str):
        return None
    m = _WZORZEC_ARKUSZA_W_FORMULE.match(formula_text)
    if not m:
        return None
    return m.group(1) or m.group(2)


def _podmien_arkusz_w_formule(formula_text, nowa_nazwa):
    if not isinstance(formula_text, str):
        return formula_text
    nazwa_escaped = nowa_nazwa.replace("'", "''")
    return _WZORZEC_ARKUSZA_W_FORMULE.sub(f"='{nazwa_escaped}'!", formula_text, count=1)


def _brak_wartosci(v):
    """True gdy wartość jest pusta (None lub pusty string)."""
    return v is None or (isinstance(v, str) and v.strip() == "")


_WZORZEC_POWTORKI_NAZWY = re.compile(r'\s\(\d+\)$')


def _to_jest_powtorka_kolizji(nazwa):
    """
    True, gdy nazwa zakladki ma sufiks ' (N)' dodany przez _unikalne_nazwy_zakladek
    z powodu kolizji nazw — co w praktyce oznacza powtorzony pomiar histerezy
    (ten sam nominalny punkt temp/RH zmierzony drugi raz).
    """
    return bool(nazwa) and bool(_WZORZEC_POWTORKI_NAZWY.search(nazwa))


def _rh_z_nazwy_zakladki(nazwa):
    """
    Parsuje NOMINALNA wilgotnosc z nazwy zakladki 'temp, rh' (np. '25, 51' -> 51.0,
    '10, 52 (2)' -> 52.0). Zwraca None gdy brak RH ('5, -') lub nie da sie sparsowac.
    Nazwa zakladki to najpewniejsze zrodlo nominalu RH — komorki F15:F19 (odczyt
    wilgotnosci) bywaja puste w zapisanych kopiach, wiec nie nadaja sie do wyboru
    reprezentanta 'najblizszego 50%'.
    """
    if not nazwa:
        return None
    baza = _WZORZEC_POWTORKI_NAZWY.sub('', str(nazwa)).strip()   # usun sufiks ' (N)'
    czesci = baza.split(',')
    if len(czesci) < 2:
        return None
    try:
        return float(czesci[1].strip().replace(',', '.'))
    except ValueError:
        return None


def _wybierz_reprezentantow_temp(wpisy, tol_temp=1.0):
    """
    wpisy: lista (temp_zaokr, rh_lub_None, czy_powtorka, payload).

    Laczymy TYLKO zakladki Z WILGOTNOSCIA (wzorcowanie RH): kilka zakladek o tej samej
    (zblizonej, <= tol_temp °C) temperaturze, ale roznej wilgotnosci, to jeden punkt
    temperaturowy → jeden reprezentant (do tabel 'tylko temperatura': Excel O1:R9,
    Word bez kolumn RH). Z grupy preferowane sa zakladki bez sufiksu powtorki ' (N)',
    a z nich ta o wilgotnosci NAJBLIZSZEJ 50%.

    Zakladki BEZ wilgotnosci ('-', czyli wzorcowanie samej temperatury) to ODDZIELNE
    punkty i NIGDY nie sa laczone — kazda dostaje wlasny wiersz. Bez tego sasiednie
    punkty temp co ~1-2 °C (np. 41,6 i 42,3 -> zaokr. 42 i 42) bledy sie sklejaly i
    czesc punktow ginela w arkuszu Wyniki.
    Zwraca liste wybranych 'payload' (bez okreslonej kolejnosci).
    """
    # Klastrowanie po temperaturze z tolerancja (kotwica = najnizsza temp w grupie).
    # Grupa moze "wchlaniac" kolejne wpisy TYLKO gdy i ona, i nowy wpis maja wilgotnosc.
    posortowane = sorted(wpisy, key=lambda x: x[0])
    grupy = []   # lista [kotwica_temp, [wpisy...], czy_grupa_z_wilgotnoscia]
    for wp in posortowane:
        ma_rh = wp[1] is not None
        if (grupy and ma_rh and grupy[-1][2]
                and abs(wp[0] - grupy[-1][0]) <= tol_temp):
            grupy[-1][1].append(wp)
        else:
            grupy.append([wp[0], [wp], ma_rh])

    wynik = []
    for _kotwica, grupa, _hum in grupy:
        kandydaci = [(rh, powtorka, payload) for (_t, rh, powtorka, payload) in grupa]
        bez_powtorki = [k for k in kandydaci if not k[1]]
        pool = bez_powtorki if bez_powtorki else kandydaci
        z_rh = [k for k in pool if k[0] is not None]
        wybrany = min(z_rh, key=lambda k: abs(k[0] - 50)) if z_rh else pool[0]
        wynik.append(wybrany[2])
    return wynik


def _znajdz_komorke_histerezy(ws_w, extra_rows=0):
    """
    Znajduje komorke z formula histerezy w arkuszu Wyniki (kolumna J, jedyna z 'E219').
    W szablonie jest to J23, ale tabela Wyniki moze dodac wiersze (>6 zakladek), przez
    co formula zjezdza w dol o `extra_rows` (np. do J26). Najpierw celujemy w
    J(23+extra_rows), a gdy tam nie ma formuly — skanujemy kolumne J.
    Zwraca (obiekt_komorki, tekst_formuly) albo (None, None).
    """
    r0 = 23 + int(extra_rows or 0)
    try:
        f = ws_w.range(f"J{r0}").formula
        if isinstance(f, str) and "E219" in f:
            return ws_w.range(f"J{r0}"), f
    except Exception:
        pass
    try:
        kol = ws_w.range("J1:J80").formula
    except Exception:
        return None, None
    for idx, val in enumerate(kol, start=1):
        f = val[0] if isinstance(val, (list, tuple)) else val
        if isinstance(f, str) and f.startswith("=") and "E219" in f:
            return ws_w.range(f"J{idx}"), f
    return None, None


def _aktualizuj_formule_histerezy(ws_w, working_final, extra_rows=0):
    """
    Ustawia komorke histerezy w arkuszu Wyniki (kolumna J; w szablonie J23,
    po dodaniu wierszy tabeli przesunieta o `extra_rows`).

    Regula:
      • jest punkt histerezy (arkusz z sufiksem ' (N)', np. '20, 44 (2)') →
        wpisujemy formule =ABS(VALUE('<baza>'!E219)-VALUE('<baza> (2)'!E219))
        (baza = nazwa arkusza bez ' (N)', np. '20, 44'); Excel pokaze ja jako
        MODUŁ.LICZBY(WARTOŚĆ(...)-WARTOŚĆ(...)) w polskim interfejsie;
      • brak powtorki (2) → w komorce wpisujemy liczbe 1 (zamiast formuly).
    """
    cell, _ = _znajdz_komorke_histerezy(ws_w, extra_rows)
    if cell is None:
        print("  [Wyniki] Nie znaleziono komorki formuly histerezy (J*, E219) — pomijam.")
        return

    rep = next((s for s in working_final if _to_jest_powtorka_kolizji(s)), None)
    if rep is None:
        try:
            cell.value = 1
            print(f"  [Wyniki] Histereza {cell.address}: brak punktu (2) → wpisano 1")
        except Exception as e:
            print(f"  [Wyniki] Nie udalo sie wpisac 1 do {cell.address}: {e}")
        return

    base = _WZORZEC_POWTORKI_NAZWY.sub('', rep).strip()
    nowa = f"=ABS(VALUE('{base}'!E219)-VALUE('{rep}'!E219))"
    try:
        cell.formula = nowa
        print(f"  [Wyniki] Histereza {cell.address}: '{base}' / '{rep}'")
    except Exception as e:
        print(f"  [Wyniki] Nie udalo sie zapisac formuly histerezy: {e}")


def _uporzadkuj_tabele_wyniki(app, ws_w, working_final):
    """
    Porzadkuje dwie tabele wynikowe w arkuszu Wyniki: B1:I9 (temp+wilgotnosc)
    i O1:R9 (tylko temp). Usuwa/czysci wiersze z #REF!, sortuje rosnaco po
    temperaturze (kol. B) z tiebreakiem po wilgotnosci (kol. C, nizsza = "wyzsza",
    wiec umieszczana nizej), dodaje wiersze pod tabela gdy len(working_final) > 6
    (klonujac formuly wzorcowego wiersza i podmieniajac nazwe arkusza). Nie zmienia
    formul w obrebie wiersza — tylko przenosi cale wiersze.

    Tabela B1:I9 (z wilgotnoscia) zawiera WYLACZNIE zakladki, ktore faktycznie
    posiadaja wilgotnosc (kol. C — wartosc odniesienia RH — daje liczbe, nie blad
    formuly); zakladki bez wilgotnosci sa z niej wykluczone (kolumny czyszczone).
    Tabela O1:R9 (tylko temp) zawiera WSZYSTKIE zakladki, ale zdeduplikowane po
    zaokraglonej temperaturze: gdy kilka zakladek dzieli ta sama temperature
    (rozne cele wilgotnosci przy wzorcowaniu RH), wybierany jest jeden reprezentant
    — patrz _wybierz_reprezentantow_temp.
    Zwraca liczbe dodanych wierszy (przesuniecie dla F24/C28/C32/E28/E32).
    """
    FIRST_ROW = 4
    LAST_ROW_DEFAULT = 9
    last_row = LAST_ROW_DEFAULT

    wiersze = []
    for r in range(FIRST_ROW, last_row + 1):
        f_bi = _czytaj_wiersz_formul(ws_w, r, _COLS_BI)
        f_or = _czytaj_wiersz_formul(ws_w, r, _COLS_OR)
        blad = any(isinstance(x, str) and "#REF" in x.upper() for x in f_bi + f_or)
        wiersze.append({
            "formula_bi": f_bi, "formula_or": f_or,
            "val_b": ws_w.range(f"B{r}").value,
            "val_c": ws_w.range(f"C{r}").value,
            "blad": blad,
        })

    valid = [w for w in wiersze if not w["blad"]]
    referenced = set()
    for w in valid:
        s = _arkusz_z_formuly(w["formula_bi"][0])
        if s:
            referenced.add(s)

    missing = [s for s in working_final if s not in referenced] if working_final else []

    if missing and valid:
        wzor = valid[0]
        nowe = []
        for nazwa in missing:
            r_nowy = last_row + 1
            ws_w.api.Rows(r_nowy).Insert()
            last_row = r_nowy
            nowy_bi = [_podmien_arkusz_w_formule(x, nazwa) for x in wzor["formula_bi"]]
            nowy_or = [_podmien_arkusz_w_formule(x, nazwa) for x in wzor["formula_or"]]
            r = last_row
            _zapisz_wiersz_formul(ws_w, r, _COLS_BI, nowy_bi)
            _zapisz_wiersz_formul(ws_w, r, _COLS_OR, nowy_or)
            nowe.append({"formula_bi": nowy_bi, "formula_or": nowy_or, "val_b": None, "val_c": None, "blad": False})
        app.calculate()
        for i, w in enumerate(nowe):
            r = LAST_ROW_DEFAULT + 1 + i
            w["val_b"] = ws_w.range(f"B{r}").value
            w["val_c"] = ws_w.range(f"C{r}").value
        wiersze.extend(nowe)
        valid = [w for w in wiersze if not w["blad"]]
    elif missing and not valid:
        # Brak wzorca do sklonowania — buduj formuly od zera wg stalego schematu.
        # BI: ='SHEET'!D239..K239 (8 kol.), OR: ='SHEET'!D246..G246 (4 kol.)
        BI_KOMORKI = ["D239", "E239", "F239", "G239", "H239", "I239", "J239", "K239"]
        OR_KOMORKI = ["D246", "E246", "F246", "G246"]
        nowe = []
        for nazwa in missing:
            r_nowy = last_row + 1
            ws_w.api.Rows(r_nowy).Insert()
            last_row = r_nowy
            nazwa_esc = nazwa.replace("'", "''")
            nowy_bi = [f"='{nazwa_esc}'!{k}" for k in BI_KOMORKI]
            nowy_or = [f"='{nazwa_esc}'!{k}" for k in OR_KOMORKI]
            _zapisz_wiersz_formul(ws_w, r_nowy, _COLS_BI, nowy_bi)
            _zapisz_wiersz_formul(ws_w, r_nowy, _COLS_OR, nowy_or)
            nowe.append({"formula_bi": nowy_bi, "formula_or": nowy_or,
                         "val_b": None, "val_c": None, "blad": False})
            print(f"    [Wyniki] Dodano wiersz z formulami dla arkusza '{nazwa}' (fallback).")
        app.calculate()
        for i, w in enumerate(nowe):
            r = LAST_ROW_DEFAULT + 1 + i
            w["val_b"] = ws_w.range(f"B{r}").value
            w["val_c"] = ws_w.range(f"C{r}").value
        wiersze.extend(nowe)
        valid = [w for w in wiersze if not w["blad"]]

    n_rows = last_row - FIRST_ROW + 1

    def _czy_liczba(v):
        return isinstance(v, (int, float)) and not isinstance(v, bool)

    def _temp_zaokr(w):
        try:
            return _round_half_away_from_zero(float(w["val_b"]))
        except (TypeError, ValueError):
            return None

    # Sprawdzaj obecnosc wilgotnosci per zakladka przez F15:F19 — to samo zrodlo,
    # ktorego uzywa _odczytaj_kalibracje_xlwings. Nie polegamy na E239 (val_c
    # z Wyniki), bo formula E239 moze dac 0 lub byc niesprawdzona w momencie
    # odczytu. Przy okazji wyciagamy nominal RH (pierwszy niepusty F15) do
    # uzycia jako klucz sortowania i dedup — jest to wartosc wpisana przez
    # skrypt, wiec zawsze poprawna.
    wb = ws_w.book
    sheet_names = {s.name for s in wb.sheets}
    for w in valid:
        arkusz = _arkusz_z_formuly(w["formula_bi"][0])
        w["arkusz"] = arkusz
        if arkusz and arkusz in sheet_names:
            rhm_vals = wb.sheets[arkusz].range("F15:F19").value
            if not isinstance(rhm_vals, list):
                rhm_vals = [rhm_vals]
            non_puste = [v for v in rhm_vals if not _brak_wartosci(v)]
            w["ma_rh"] = bool(non_puste)
            try:
                w["rh_ref"] = float(non_puste[0]) if non_puste else None
            except (TypeError, ValueError):
                w["rh_ref"] = None
        else:
            w["ma_rh"] = _czy_liczba(w["val_c"])
            try:
                w["rh_ref"] = float(w["val_c"]) if w["ma_rh"] else None
            except (TypeError, ValueError):
                w["rh_ref"] = None
        w["powtorka"] = _to_jest_powtorka_kolizji(arkusz)

    def _klucz_bi(w):
        temp = w["val_b"]
        try:
            temp = float(temp)
        except (TypeError, ValueError):
            return (float('inf'), 1, float('inf'), float('inf'))
        rounded = _round_half_away_from_zero(temp)
        rh_ref = w.get("rh_ref")
        rh_val = float(rh_ref) if rh_ref is not None else 0.0
        # Powtorki histerezy (sufiks ' (N)') na koncu grupy temperaturowej;
        # w ramach tej samej grupy RH rosnie (8 → 49 → 82 → 49(2)).
        is_dup = 1 if w.get("powtorka") else 0
        return (rounded, is_dup, rh_val, temp)

    valid_bi_sorted = sorted((w for w in valid if w["ma_rh"]), key=_klucz_bi)

    wpisy_or = [
        (_temp_zaokr(w), _rh_z_nazwy_zakladki(w.get("arkusz")), w["powtorka"], w)
        for w in valid if _temp_zaokr(w) is not None
    ]
    reprezentanci_or = _wybierz_reprezentantow_temp(wpisy_or)
    reprezentanci_or_sorted = sorted(reprezentanci_or, key=_temp_zaokr)

    for i, w in enumerate(valid_bi_sorted):
        r = FIRST_ROW + i
        _zapisz_wiersz_formul(ws_w, r, _COLS_BI, w["formula_bi"])
    for i in range(len(valid_bi_sorted), n_rows):
        r = FIRST_ROW + i
        _wyczysc_wiersz(ws_w, r, _COLS_BI)

    for i, w in enumerate(reprezentanci_or_sorted):
        r = FIRST_ROW + i
        _zapisz_wiersz_formul(ws_w, r, _COLS_OR, w["formula_or"])
    for i in range(len(reprezentanci_or_sorted), n_rows):
        r = FIRST_ROW + i
        _wyczysc_wiersz(ws_w, r, _COLS_OR)

    return last_row - LAST_ROW_DEFAULT


def _dostosuj_xlwings(app, sciezka_pliku, dane_zakladek, dane_ef_kopia, rekord, nowa_nazwa, ark_wyniki, f24_val, sciezki_linkowane=None, _pierwsza_kopia=True, _cache_fg=None):
    """
    Otwiera kopie xlsx przez xlwings (COM Excel):
      1. Dopasowuje liczbe zakladek roboczych do len(dane_zakladek).
      2. Zmienia nazwy zakladek (przez tymczasowe, zeby uniknac konfliktow).
      3. Wypelnia C15:C19 / D15:D19 / E15:E19 / F15:F19 (dane z Strona 3).
      4. Wypelnia komorki naglowkowe i stopkowe w zakladkach roboczych (Etap 5).
      5. Wypelnia F24, C28, C32, E28:G28, E32:G32 w arkuszu Wyniki (Etap 6).
      6. Wypelnia F/G warunki srodowiskowe w Strona 3, oblicza zakresy przez Wzory.xls.
      7. Zapisuje plik.
    Uzycie COM Excel gwarantuje zachowanie wszystkich formul w arkuszu.
    Zwraca: slownik {"temp_min", "temp_max", "wilg_min", "wilg_max"} lub None.
    """
    # KLUCZOWE: przed otwarciem przepisujemy linki w XML na LOKALNE pliki.
    # Inaczej kopia wskazuje na \\plum4\... — Excel przy Open() szuka tego pliku
    # w sieci, a gdy serwer jest niedostepny SMB-lookup zawiesza sie i po kilku
    # kopiach RPC Excela pada (-2147023174). Lokalne pliki istnieja i sa otwarte
    # w tej sesji → Excel resolwuje natychmiast, bez dotykania serwera.
    # Linki serwerowe przywracamy na koncu (finally w utworz_kopie).
    if sciezki_linkowane:
        _przywroc_linki_w_xml(sciezka_pliku, sciezki_linkowane, cicho=True)

    # EnableEvents=False: blokuje Application.WorkbookOpen/SheetActivate/
    # BeforeSave/BeforeClose w Obliczenia.xls/Wzory.xls → \\plum4 → crash.
    # xlManual: na czas Open() ORAZ calej fazy STRUKTURALNEJ (kopiowanie zakladek,
    # zmiana nazw, wypelnianie komorek). Przy wielu punktach (np. 34 zakladki) kazde
    # kopiowanie/zapis w trybie automatycznym wymuszalo pelne przeliczenie rosnacego
    # skoroszytu z linkami zewnetrznymi — stad zawieszanie i pad Excela
    # (OLE 0x800a01a8 'Object required'). xlAutomatic przywracamy dopiero przed
    # etapem obliczen; miejsca, ktore potrzebuja wynikow, i tak wolaja Calculate().
    _t0 = time.time()
    app.api.EnableEvents = False
    app.api.Calculation = -4135  # xlCalculationManual
    _log_etap(f"otwieram kopie: {os.path.basename(sciezka_pliku)}", _t0)
    wb = _open_book_hidden(app, sciezka_pliku, update_links=False)
    _log_etap("otwarta (tryb przeliczania: RECZNY)", _t0)
    # Przekieruj linki zewnetrzne na lokalnie otwarte pliki (Obliczenia, Wzory).
    # Bez tego formuly lancuchujace do =[Obliczenia]!te_6(...) zwracaja None/blad
    # nawet gdy Obliczenia jest otwarte, bo Excel nie moze dopasowac zapisanej
    # sciezki serwerowej do lokalnego pliku — val_b w Wyniki byloby None, sortowanie
    # i zamrazanie wartosci w Wyniki nie dzialaloby poprawnie.
    # UWAGA: ChangeLink wywolujemy TYLKO dla znanych plikow z sciezki_linkowane.
    # UpdateLink (dla nieznanych) omijamy — moze powodowac nieoczekiwane otwarcia
    # plikow serwera lub zapis kopii do blednej sciezki.
    if sciezki_linkowane:
        try:
            sources = wb.api.LinkSources(1)  # 1 = xlLinkTypeExcelLinks
            if sources:
                for _src in sources:
                    _nazwa = _nazwa_pliku_z_linku(_src)
                    if not _nazwa:
                        continue
                    _nowa = sciezki_linkowane.get(_nazwa.lower())
                    if _nowa:
                        try:
                            wb.api.ChangeLink(Name=_src, NewName=_nowa, Type=1)
                            if _pierwsza_kopia:
                                print(f"    [ChangeLink] {_nazwa}: lokalny plik OK")
                        except Exception as _ce:
                            _warn(f"ChangeLink nie powiodl sie dla: {_nazwa}\n"
                                  f"        stary link: {_src}\n"
                                  f"        nowy  link: {_nowa}\n"
                                  f"        Blad: {type(_ce).__name__}: {_ce}", indent="    ")
                    else:
                        _warn(f"Link '{_nazwa}' — brak lokalnej kopii w folderze roboczym\n"
                              f"        Sciezka serwerowa: {_info_serwer(_nazwa)}\n"
                              f"        Pelny link w pliku: {_src}", indent="    ")
        except Exception as _le:
            _warn(f"LinkSources() — nie mozna odczytac linkow z pliku\n"
                  f"      {type(_le).__name__}: {_le}")
    try:
        # arkusz zbiorczy nie jest zakladka punktu — nie usuwamy go ani nie przemianowujemy
        wykluczone = {ark_wyniki}
        working = [s.name for s in wb.sheets if s.name not in wykluczone]
        N = len(dane_zakladek)

        if not working and N > 0:
            print("  [BLAD] Brak zakladek roboczych — pomijam.")
            return

        if N == 0:
            # Dla kopii bez aktywnych blokow usuwamy wszystkie zakladki robocze.
            wszystkie_nazwy = {s.name for s in wb.sheets}
            if ark_wyniki in wszystkie_nazwy:
                for ws_name in list(working):
                    wb.sheets[ws_name].delete()
            else:
                # Gdy brak arkusza chronionego, nie usuwamy ostatniej zakladki,
                # bo Excel nie pozwoli zapisac skoroszytu bez zadnego arkusza.
                for ws_name in list(working[:-1]):
                    wb.sheets[ws_name].delete()
            working_final = [s.name for s in wb.sheets if s.name not in wykluczone]
        else:
            first_ws_name = working[0]

            # --- Dopasuj liczbe zakladek ---
            if len(working) > N:
                for ws_name in working[N:]:
                    wb.sheets[ws_name].delete()
                working = working[:N]

            elif len(working) < N:
                do_dodania = N - len(working)
                _log_etap(f"kopiuje zakladki: {len(working)} -> {N} "
                          f"({do_dodania} operacji kopiowania)", _t0)
                for _i in range(do_dodania):
                    _t_kopii = time.time()
                    before = {s.name for s in wb.sheets}
                    src = wb.sheets[first_ws_name]
                    if ark_wyniki in {s.name for s in wb.sheets}:
                        src.api.Copy(Before=wb.sheets[ark_wyniki].api)
                    else:
                        src.api.Copy(After=wb.sheets[-1].api)
                    after = {s.name for s in wb.sheets}
                    new_name = (after - before).pop()
                    working.append(new_name)
                    if _i == 0 and do_dodania >= PROG_OSTRZEZENIA_KOPII:
                        # Czas MIERZYMY, nie zakladamy — zalezy od maszyny i od liczby
                        # obiektow OLE (Equation/Word) w zakladce szablonu.
                        _sek = time.time() - _t_kopii
                        _warn(f"Do wykonania {do_dodania} kopii zakladek (punktow: {N}).\n"
                              f"      Pierwsza trwala {_sek:.1f} s -> szacunkowo "
                              f"~{do_dodania * _sek / 60:.0f} min NA PLIK.\n"
                              f"      Kopiowanie jest wolne, bo zakladka szablonu zawiera obiekty\n"
                              f"      OLE (Equation/Word) — Excel duplikuje kazdy z nich i uruchamia\n"
                              f"      do tego Worda. Przy wielu punktach Word potrafi zglosic\n"
                              f"      'Za duzo otwartych plikow', a Excel padnac. Trwale rozwiazanie:\n"
                              f"      zamiana tych obiektow w szablonie na obrazki.")
                    if (_i + 1) % 5 == 0 or _i + 1 == do_dodania:
                        _log_etap(f"   skopiowano {_i + 1}/{do_dodania} zakladek", _t0)

            working = [s.name for s in wb.sheets if s.name not in wykluczone]

            # --- Zmien nazwy przez tymczasowe (unika konfliktow) ---
            _log_etap(f"zmieniam nazwy {len(working)} zakladek", _t0)
            for i, ws_name in enumerate(working):
                wb.sheets[ws_name].name = f"__tmp_{i}__"

            docelowe_nazwy = _unikalne_nazwy_zakladek(dane_zakladek)
            working_tmp = [s.name for s in wb.sheets if s.name not in wykluczone]
            for i, ws_name in enumerate(working_tmp):
                wb.sheets[ws_name].name = docelowe_nazwy[i]
            _log_etap(f"nazwy ustawione: {', '.join(docelowe_nazwy[:4])}"
                      f"{' ...' if len(docelowe_nazwy) > 4 else ''}", _t0)

            # --- Wypelnij komorki ---
            working_final = [s.name for s in wb.sheets if s.name not in wykluczone]

        prefiks, rok, nr_fab, _ = _parsuj_nazwe_pliku(nowa_nazwa)
        # SAMA data, bez godziny — wartosc trafia do komorek dat przy podpisach
        # (B228/H228 oraz C28/C32 w 'Wyniki'). datetime.now() zapisywalby tam czas,
        # ktory byl widoczny na pasku formuly ('06.08.2026 10:14:42').
        dzis = datetime.date.today()
        parametry_cc04 = _parametry_typu_cc04(rekord)
        if rekord.get("IS_CC04_PROTO") and parametry_cc04 is None:
            raw = _cell_to_str(rekord.get("CC04_RAW"))
            print(f"  [UWAGA] Nieznany typ CC-04 dla kopii '{nowa_nazwa}': '{raw}'.")

        # Higrometr do K18 zalezy od komory (CC / CC-04) — patrz konfiguracja na gorze pliku.
        _k18_wartosc = _higrometr_k18(bool(rekord.get("IS_CC04_PROTO")))
        _log_etap(f"K18 (higrometr) dla komory "
                  f"{'CC-04' if rekord.get('IS_CC04_PROTO') else 'CC'}: '{_k18_wartosc}'", _t0)

        _log_etap(f"wypelniam komorki w {len(working_final)} zakladkach...", _t0)
        for i, ws_name in enumerate(working_final):
            if (i + 1) % 5 == 0 or i == 0 or i + 1 == len(working_final):
                _log_etap(f"   zakladka {i + 1}/{len(working_final)}: '{ws_name}'", _t0)
            ws = wb.sheets[ws_name]
            zd = dane_zakladek[i]
            ef = dane_ef_kopia[i] if i < len(dane_ef_kopia) else None

            for offset in range(5):
                val = None
                fmt = None
                if isinstance(ef, dict) and "C_dane" in ef:
                    c_dane = ef.get("C_dane") or []
                    if offset < len(c_dane):
                        val = c_dane[offset]
                        c_fmt = ef.get("C_fmt") or []
                        fmt = c_fmt[offset] if offset < len(c_fmt) else None
                elif val is None:
                    val = zd["L_dane"][offset]
                    l_fmt = zd.get("L_fmt") or []
                    fmt = l_fmt[offset] if offset < len(l_fmt) else None
                if val is not None:
                    cell = ws.cells(15 + offset, 3)
                    cell.value = val
                    if fmt:
                        cell.api.NumberFormat = fmt

            for offset in range(5):
                val = None
                fmt = None
                if isinstance(ef, dict) and "D_dane" in ef:
                    d_dane = ef.get("D_dane") or []
                    if offset < len(d_dane):
                        val = d_dane[offset]
                        d_fmt = ef.get("D_fmt") or []
                        fmt = d_fmt[offset] if offset < len(d_fmt) else None
                elif val is None:
                    val = zd["M_dane"][offset]
                    m_fmt = zd.get("M_fmt") or []
                    fmt = m_fmt[offset] if offset < len(m_fmt) else None
                if val is not None:
                    cell = ws.cells(15 + offset, 4)
                    cell.value = val
                    if fmt:
                        cell.api.NumberFormat = fmt

            if isinstance(ef, dict):
                e_dane = ef.get("E_dane") or []
                e_fmt_list = ef.get("E_fmt") or []
                f_dane = ef.get("F_dane") or []
                f_fmt_list = ef.get("F_fmt") or []
                for offset in range(5):
                    val = e_dane[offset] if offset < len(e_dane) else None
                    fmt = e_fmt_list[offset] if offset < len(e_fmt_list) else None
                    if val is not None:
                        cell = ws.cells(15 + offset, 5)
                        cell.value = val
                        if fmt:
                            cell.api.NumberFormat = fmt
                    if fmt:
                        ws.cells(15 + offset, 7).api.NumberFormat = fmt  # G15:G19
                for offset in range(5):
                    val = f_dane[offset] if offset < len(f_dane) else None
                    fmt = f_fmt_list[offset] if offset < len(f_fmt_list) else None
                    if val is not None:
                        cell = ws.cells(15 + offset, 6)
                        cell.value = val
                        if fmt:
                            cell.api.NumberFormat = fmt
                    if fmt:
                        ws.cells(15 + offset, 8).api.NumberFormat = fmt  # H15:H19

            # --- Etap 5: naglowki i stopki ---
            ws.range("E4").value = f"{prefiks}/LA/TH/{rok}"
            ws.range("G6").value = nr_fab
            k4 = zd.get("K4_val")
            if k4 is not None:
                ws.range("K4").value = k4
            if rekord.get("B") is not None:
                ws.range("E5").value = rekord["B"]
            if rekord.get("D") is not None:
                ws.range("E6").value = rekord["D"]
            if rekord.get("K") is not None:
                ws.range("H57").value = rekord["K"]

            # Dla protokolow CC-04 wpisz stale parametry przyrzadu.
            if parametry_cc04 is not None:
                ws.range("K11").value = parametry_cc04["K11"]
                ws.range("K12").value = parametry_cc04["K12"]
                ws.range("K13").value = parametry_cc04["K13"]
                ws.range("K17").value = parametry_cc04["K17"]

            # K18 — higrometr punktu rosy: wartosc wg KOMORY (HIGROMETR_K18_WG_KOMORY),
            # ale dla punktow TYLKO-TEMPERATURA (brak RH w nazwie zakladki) zawsze "-".
            _rh_nom = _rh_z_nazwy_zakladki(zd.get("nazwa"))
            ws.range("K18").value = "-" if _rh_nom is None else _k18_wartosc

            o25 = zd.get("o25_val")
            if o25 is not None:
                try:
                    ws.range("O25").value = o25
                except Exception:
                    pass

            ws.range("B228").value = dzis
            ws.range("H228").value = dzis
            ws.range("B230").value = PODPISUJACY_1
            ws.range("H230").value = PODPISUJACY_2

        # Faza STRUKTURALNA skonczona — wracamy do trybu automatycznego, zeby UDF-y
        # z Wzory.xls/Obliczenia liczyly sie tak jak dotad podczas Calculate() ponizej.
        try:
            app.api.Calculation = -4105  # xlCalculationAutomatic
        except Exception:
            pass
        app.api.DisplayAlerts = False
        app.api.AskToUpdateLinks = False
        app.api.EnableEvents = False
        _log_etap(f"przeliczam {len(wb.sheets)} zakladek (1. przebieg)...", _t0)
        for _ws in wb.sheets:
            _ws.api.Calculate()   # przelicz tylko kopie (nie Obliczenia/Wzory) — blokuje Worksheet_Calculate w .xls
        app.api.EnableEvents = True
        _log_etap("przeliczono", _t0)

        # --- Etap 6: arkusz Wyniki ---
        extra_rows = 0
        if ark_wyniki in {s.name for s in wb.sheets}:
            _log_etap(f"arkusz '{ark_wyniki}': porzadkuje tabele wynikow...", _t0)
            ws_w = wb.sheets[ark_wyniki]
            extra_rows = _uporzadkuj_tabele_wyniki(app, ws_w, working_final)
            _aktualizuj_formule_histerezy(ws_w, working_final, extra_rows)
            if f24_val is not None:
                ws_w.range(f"F{24 + extra_rows}").value = f24_val
            ws_w.range(f"C{28 + extra_rows}").value = dzis
            ws_w.range(f"C{32 + extra_rows}").value = dzis
            ws_w.range(f"E{28 + extra_rows}").value = PODPISUJACY_1   # scalone E28:G28
            ws_w.range(f"E{32 + extra_rows}").value = PODPISUJACY_2   # scalone E32:G32

        # --- Etap 7: warunki środowiskowe (Strona 3 F/G + Wzory.xls) ---
        _log_etap("licze warunki srodowiskowe (Pom. nr 9 + Wzory.xls)...", _t0)
        zakresy_srodowiskowe = _oblicz_warunki_srodowiskowe(
            app, wb, dane_zakladek, NR_POMIESZCZENIA, MODEL_CZUJNIKA, _cache_fg=_cache_fg
        )
        _log_etap("warunki srodowiskowe policzone", _t0)

        # Przelicz i zapisz — wb.save() utrwala obliczone wartosci jako cache Excela.
        # Dzieki temu po otwarciu kopii bez plikow linkowanych (Obliczenia, Wzory)
        # i odrzuceniu aktualizacji linkow Excel pokazuje te zakeszowane wartosci.
        app.api.DisplayAlerts = False
        app.api.AskToUpdateLinks = False
        _log_etap("przeliczam przed zapisem (2. przebieg)...", _t0)
        for _ws in wb.sheets:
            _ws.api.Calculate()   # przelicz tylko kopie przed zapisem
        # EnableEvents pozostaje False (ustawione przed open, przywrocone w finally)

        # Pasek zakladek: po otwarciu pliku maja byc widoczne WSZYSTKIE zakladki punktow,
        # a aktywny arkusz zbiorczy. KOLEJNOSC MA ZNACZENIE — activate() sam przewija pasek
        # do aktywnego arkusza, wiec przewijanie musi byc PO aktywacji (wczesniej bylo
        # odwrotnie i po otwarciu widac bylo tylko 'Wyniki').
        try:
            if ark_wyniki in {s.name for s in wb.sheets}:
                wb.sheets[ark_wyniki].activate()
            okno = wb.api.Windows(1)
            try:
                okno.TabRatio = TAB_RATIO      # szerszy pasek zakladek kosztem paska przewijania
            except Exception:
                pass
            okno.ScrollWorkbookTabs(Sheets=-wb.api.Sheets.Count)   # na sam poczatek
        except Exception:
            pass

        app.api.DisplayAlerts = False
        app.api.AskToUpdateLinks = False
        # EnableEvents pozostaje False — BeforeSave/AfterSave nie strzela
        _log_etap("zapisuje plik...", _t0)
        try:
            wb.save()
        except Exception:
            wb.api.Save()
        _log_etap("ZAPISANO — kopia gotowa", _t0)
    except Exception as _exc:
        # Najczestszy przypadek: Excel padl w trakcie (OLE 0x800a01a8 'Object required'
        # albo RPC -2147023174). Log mowi, na ktorym etapie i po ilu sekundach.
        _err(f"Przerwano prace nad kopia '{nowa_nazwa}' po {time.time() - _t0:.1f}s\n"
             f"      {type(_exc).__name__}: {_exc}\n"
             f"      (jesli to pad Excela — zamknij wszystkie okna Excela i uruchom ponownie)")
        raise
    finally:
        try:
            wb.close()
        except Exception:
            pass
        try:
            app.api.Calculation = -4105  # xlCalculationAutomatic
        except Exception:
            pass
        try:
            app.api.EnableEvents = True
        except Exception:
            pass
    return zakresy_srodowiskowe


def _odczytaj_kalibracje_xlwings(app, sciezka_pliku, ark_wyniki, enable_diag=False, sciezki_linkowane=None):
    """
    Czyta dane kalibracyjne z zakladek roboczych po pelnym przeliczeniu formul.
    Wykrywa obecnosc aktywnych danych wilgotnosci (RHm w F15:F19) NIEZALEZNIE
    dla kazdej zakladki — zakladki z wilgotnoscia maja czytane 8 wartosci z wiersza
    239 (t/RH/tm/RHm/Δt/ΔRH/Ut/URH), zakladki bez wilgotnosci maja czytane 4
    wartosci z wiersza 246 (t/tm/Δt/U). Pozwala to na mieszane kopie (czesc
    zakladek z wilgotnoscia, czesc bez).
    sciezki_linkowane: opcjonalny slownik {nazwa_pliku.lower(): pelna_lokalna_sciezka}
    plikow linkowanych (PLIKI_LINKOWANE) aktualnie otwartych w tej samej sesji
    Excela — uzywany do jawnego przekierowania linku (ChangeLink) na lokalnie
    otwarty plik, niezalegle od tego, jaka (czasem zniekszalcona/nieaktualna)
    sciezke ma zapisana plik na dysku.
    Zwraca: (kalibracja, klasa_wilgotnosci), gdzie klasa_wilgotnosci to jedna z:
      "brak"     — zadna zakladka nie ma wilgotnosci (szablon Word: tylko temp)
      "pelna"    — wszystkie zakladki maja wilgotnosc (szablon Word: zakres)
      "mieszana" — czesc zakladek ma wilgotnosc, czesc nie (szablon Word: zakres + temp)
    Kazdy wpis 'kalibracja' zawiera dodatkowo "nazwa" (nazwa zakladki) — uzywane
    przy wyborze reprezentantow tej samej temperatury w tabeli 'tylko temp'.
    """
    # Przed otwarciem przepisujemy linki na LOKALNE pliki — tak samo jak w
    # _dostosuj_xlwings. Zapobiega szukaniu \\plum4 w sieci podczas Open()
    # (niedostepny serwer => zawieszenie SMB => RPC crash).
    if sciezki_linkowane:
        _przywroc_linki_w_xml(sciezka_pliku, sciezki_linkowane, cicho=True)

    # Application.WorkbookOpen/BeforeClose w Obliczenia.xls strzela przy kazdej
    # operacji na dowolnym workbooku i probuje dostac sie do \\plum4 → RPC crash.
    # xlManual tylko na czas Open() — po otwarciu natychmiast przywracamy xlAutomatic
    # zeby UDF dzialaly normalnie podczas pozniejszego Calculate().
    app.api.EnableEvents = False
    app.api.Calculation = -4135  # xlCalculationManual — tylko na czas Open()
    wb = _open_book_hidden(app, sciezka_pliku, update_links=False)
    try:
        app.api.Calculation = -4105  # przywroc xlAutomatic natychmiast po otwarciu
    except Exception:
        pass
    kalibracja = []
    try:
        # Przekieruj linki zewnetrzne na lokalnie otwarte pliki PLIKI_LINKOWANE
        # (po nazwie pliku, niezaleznie od sciezki zapisanej w pliku — Excel
        # zapisuje spacje w Target jako %20, co psuje dopasowanie po prostej
        # sciezce, a po przywroceniu sciezek serwerowych ten link i tak nie
        # jest dostepny z tej maszyny). Bez tego CalculateFullRebuild ponizej
        # moze pozostawic powiazane komorki puste/None.
        try:
            sources = wb.api.LinkSources(1)  # 1 = xlLinkTypeExcelLinks
            if sources:
                for src in sources:
                    nowa_sciezka = None
                    nazwa = None
                    if sciezki_linkowane:
                        nazwa = _nazwa_pliku_z_linku(src)
                        if nazwa:
                            nowa_sciezka = sciezki_linkowane.get(nazwa.lower())
                    try:
                        if nowa_sciezka:
                            wb.api.ChangeLink(Name=src, NewName=nowa_sciezka, Type=1)
                            if enable_diag:
                                print(f"    [ChangeLink] {nazwa}: lokalny plik OK")
                        else:
                            # Brak lokalnej kopii — pomijamy (serwer niedostepny,
                            # wartosci z _dostosuj_xlwings sa juz w cache pliku).
                            _warn(f"Link '{nazwa or src}' — brak lokalnej kopii\n"
                                  f"        Sciezka serwerowa: {_info_serwer(nazwa or '')}\n"
                                  f"        Pelny link w pliku: {src}\n"
                                  f"        (pomijamy — cache z _dostosuj_xlwings wystarczy)", indent="    ")
                    except Exception as exc:
                        _warn(f"ChangeLink nie powiodl sie dla: {nazwa or src}\n"
                              f"        Nowa sciezka: {nowa_sciezka}\n"
                              f"        Blad: {type(exc).__name__}: {exc}", indent="    ")
        except Exception as exc:
            _warn(f"LinkSources() — nie mozna odczytac linkow z pliku\n"
                  f"      {type(exc).__name__}: {exc}")

        # arkusz zbiorczy nie jest zakladka punktu — nie usuwamy go ani nie przemianowujemy
        wykluczone = {ark_wyniki}

        app.api.DisplayAlerts = False
        app.api.AskToUpdateLinks = False
        for _ws in wb.sheets:
            _ws.api.Calculate()
        # EnableEvents pozostaje False (ustawione przed open, przywrocone w finally)

        def _brak_wartosci(v):
            return v is None or (isinstance(v, str) and v.strip() == "")

        working_final_list = [s.name for s in wb.sheets if s.name not in wykluczone]

        # --- Wykryj obecnosc aktywnych danych wilgotnosci (RHm, F15:F19) per zakladka ---
        wilgotnosc_per_zakladka = {}
        for ws_name in working_final_list:
            rhm_vals = wb.sheets[ws_name].range("F15:F19").value
            if not isinstance(rhm_vals, list):
                rhm_vals = [rhm_vals]
            wilgotnosc_per_zakladka[ws_name] = any(not _brak_wartosci(v) for v in rhm_vals)

        liczba_z_wilgotnoscia = sum(1 for v in wilgotnosc_per_zakladka.values() if v)
        if liczba_z_wilgotnoscia == 0:
            klasa_wilgotnosci = "brak"
        elif liczba_z_wilgotnoscia == len(wilgotnosc_per_zakladka):
            klasa_wilgotnosci = "pelna"
        else:
            klasa_wilgotnosci = "mieszana"

        _diag_done = False
        for ws_name in working_final_list:
            ws_cal = wb.sheets[ws_name]

            if wilgotnosc_per_zakladka[ws_name]:
                d = ws_cal.range("D239").value
                e = ws_cal.range("E239").value
                f = ws_cal.range("F239").value
                g = ws_cal.range("G239").value
                h = ws_cal.range("H239").value
                ii = ws_cal.range("I239").value
                jj = ws_cal.range("J239").value
                kk = ws_cal.range("K239").value

                if enable_diag and not _diag_done and _brak_wartosci(d):
                    print(f"      [DIAG D239:K239] {ws_cal.range('D239:K239').value!r}")
                    _diag_done = True

                if any(_brak_wartosci(v) for v in (d, f, h, jj)):
                    pas = ws_cal.range("C239:N239").value
                    if isinstance(pas, list):
                        kandydaci = [v for v in pas if not _brak_wartosci(v)]
                        if len(kandydaci) >= 8:
                            d, e, f, g, h, ii, jj, kk = kandydaci[:8]

                print(f"    [Kalibracja {ws_name}] t={d!r} RH={e!r} tm={f!r} RHm={g!r} "
                      f"dt={h!r} dRH={ii!r} Ut={jj!r} URH={kk!r}")
                kalibracja.append({
                    "nazwa":          ws_name,
                    "wartosc_odn":    d,
                    "wartosc_odn_RH": e,
                    "zmierzona":      f,
                    "zmierzona_RH":   g,
                    "poprawka":       h,
                    "poprawka_RH":    ii,
                    "niepewnosc":     jj,
                    "niepewnosc_RH":  kk,
                })
            else:
                d = ws_cal.range("D246").value
                e = ws_cal.range("E246").value
                f = ws_cal.range("F246").value
                g = ws_cal.range("G246").value

                # Diagnostyka posrednich komorek (tylko pierwsza kopia, gdy enable_diag=True).
                if enable_diag and not _diag_done and (d is None or f is None or g is None):
                    print(f"      [DIAG D246] formula={ws_cal.range('D246').formula!r}  value={ws_cal.range('D246').value!r}")
                    print(f"      [DIAG F246] formula={ws_cal.range('F246').formula!r}  value={ws_cal.range('F246').value!r}")
                    print(f"      [DIAG G246] formula={ws_cal.range('G246').formula!r}  value={ws_cal.range('G246').value!r}")
                    for cell_addr in ('H57', 'G137', 'E220', 'G220', 'J220'):
                        rng = ws_cal.range(cell_addr)
                        print(f"      [DIAG {cell_addr}] formula={rng.formula!r}  value={rng.value!r}")
                    _diag_done = True

                # Fallback: szukaj 4 niepustych wartosci w szerszym pasie,
                # bo przy scaleniach dane moga wpasc nie w D:E:F:G.
                if _brak_wartosci(d) or _brak_wartosci(f) or _brak_wartosci(g):
                    pas = ws_cal.range("C246:J246").value
                    if isinstance(pas, list):
                        kandydaci = [v for v in pas if not _brak_wartosci(v)]
                        if len(kandydaci) >= 4:
                            d, e, f, g = kandydaci[:4]

                print(f"    [Kalibracja {ws_name}] REF={d!r}  ZM={e!r}  POP={f!r}  NIEP={g!r}")
                kalibracja.append({
                    "nazwa":       ws_name,
                    "wartosc_odn": d,
                    "zmierzona":   e,
                    "poprawka":    f,
                    "niepewnosc":  g,
                })
    finally:
        wb.close()
        try:
            app.api.Calculation = -4105  # xlCalculationAutomatic
        except Exception:
            pass
        try:
            app.api.EnableEvents = True
        except Exception:
            pass
    return kalibracja, klasa_wilgotnosci


# =============================================================================
# TWORZENIE KOPII
# =============================================================================

def utworz_kopie(folder, szablon_plik, dane, dane_zakladek, dane_ef, ark_wyniki, f24_per_kopia):
    """
    Dla każdego rekordu z Strona 2:
      1. Generuje nazwę kopii i kopiuje szablon (shutil.copy2).
      2. Modyfikuje kopie przez xlwings (jeden proces Excel dla wszystkich).
            3. Przywraca w kopiach sciezki linkow zewnetrznych do serwera.
    Zwraca listę nazw utworzonych kopii.
    """
    sciezka_szablonu = os.path.join(folder, szablon_plik)
    n = len(dane)
    kopie = []
    dane_zakladek_per_kopia = []
    zakresy_per_kopia = []

    # Krok 1: utwórz wszystkie kopie (szybkie kopiowanie pliku)
    for j, rekord in enumerate(dane):
        # numer przyrzadu = jego pozycja w protokole (Przyrzady wzorcowane 1, 2, 3...).
        # Bierzemy go z rekordu, bo po odfiltrowaniu przyrzadow bez pomiarow
        # indeks w petli nie odpowiada juz pozycji na Stronie 2.
        nowa_nazwa = generuj_nazwe_pliku(szablon_plik, rekord["O"], _numer_do_nazwy(rekord),
                                         nr_przyrzadu=rekord.get("_nr_przyrzadu", j + 1))
        sciezka_kopii = os.path.join(folder, _bezpieczna_nazwa_pliku(nowa_nazwa))
        if os.path.exists(sciezka_kopii):
            print(f"  [UWAGA] Plik już istnieje, nadpisuję: {nowa_nazwa}")
        shutil.copy2(sciezka_szablonu, sciezka_kopii)
        kopie.append((j, nowa_nazwa, sciezka_kopii, rekord))

    print(f"  Skopiowano {n} plików. Modyfikuję zakładki i komórki przez Excel COM...")

    # Krok 2: modyfikuj kopie — KAZDA w swiezej sesji Excela.
    # Kopiowanie zakladek z obiektami OLE zostawia w procesie Excela zasoby, ktorych on
    # nie zwalnia; przy kilku kopiach konczylo sie to padem (RPC -2147023170) w trakcie
    # drugiej kopii. Restart procesu miedzy kopiami kosztuje kilka sekund.
    app = None
    linked_wbs = []
    sciezki_linkowane = {}
    dane_kalibracji = []
    klasa_wilg_per_kopia = []
    try:
        _cache_fg = {}
        for j, nowa_nazwa, sciezka_kopii, rekord in kopie:
            print(f"[{j+1:>{len(str(n))}}/{n}] {nowa_nazwa}")
            _zamknij_sesje_excel(app, linked_wbs)
            app, linked_wbs, sciezki_linkowane = _nowa_sesja_excel(folder)
            dane_ef_kopia_surowe = dane_ef[j] if j < len(dane_ef) else []

            if USUWAJ_PUSTE_BLOKI_KOPII_S3:
                aktywne_idx, dane_zakladek_kopia, dane_ef_kopia = _wybierz_aktywne_bloki_kopii(
                    dane_zakladek,
                    dane_ef_kopia_surowe,
                )
                if aktywne_idx:
                    zakresy = _formatuj_zakresy_wierszy_s3(aktywne_idx, START_ROW_S3, BLOK_S3)
                    print(f"    Aktywne bloki Strona 3 (wiersze): {zakresy}")
                else:
                    print("    [UWAGA] Brak aktywnych blokow Strona 3 dla tej kopii — zostanie tylko arkusz Wyniki.")
            else:
                dane_zakladek_kopia = list(dane_zakladek)
                dane_ef_kopia = list(dane_ef_kopia_surowe)

            dane_zakladek_per_kopia.append(dane_zakladek_kopia)
            f24_val = f24_per_kopia[j] if j < len(f24_per_kopia) else None
            zakresy_srod = _dostosuj_xlwings(app, sciezka_kopii, dane_zakladek_kopia, dane_ef_kopia, rekord, nowa_nazwa, ark_wyniki, f24_val, sciezki_linkowane=sciezki_linkowane, _pierwsza_kopia=(j == 0), _cache_fg=_cache_fg)
            zakresy_per_kopia.append(zakresy_srod)

        print("  Odczyt kalibracji po zakonczeniu wypelniania wszystkich kopii...")
        # Swieza sesja rowniez na odczyt kalibracji (poprzednia obrobila ostatnia kopie).
        _zamknij_sesje_excel(app, linked_wbs)
        app, linked_wbs, sciezki_linkowane = _nowa_sesja_excel(folder)
        for j, nowa_nazwa, sciezka_kopii, _ in kopie:
            print(f"    [Kalibracja {j+1:>{len(str(n))}}/{n}] {nowa_nazwa}")
            kal, klasa = _odczytaj_kalibracje_xlwings(
                app, sciezka_kopii, ark_wyniki, enable_diag=(j == 0), sciezki_linkowane=sciezki_linkowane)
            dane_kalibracji.append(kal or [])
            klasa_wilg_per_kopia.append(klasa)
    finally:
        # Zamkniecie sesji odporne na pad Excela — inaczej maskowaloby oryginalny blad
        # i uniemozliwialo przywrocenie linkow ponizej.
        _zamknij_sesje_excel(app, linked_wbs)

        # Przywracanie linkow przez XML — tutaj w finally, żeby wykonało się
        # zawsze (nawet przy crash Excel podczas odczytu kalibracji).
        # ChangeLink() przez xlwings wymuszałoby przeliczenie z nową ścieżką UNC;
        # jeśli serwer jest niedostępny, Excel zapisuje błędy jako cache
        # i inne osoby nie widzą wartości po kliknięciu "Nie" w dialogu aktualizacji.
        if LINKI_SERWEROWE:
            serwery = list(LINKI_SERWEROWE.values())
            print(f"  === Przywracam linki serwerowe w {len(kopie)} plik(ach) ===")
            for s in serwery:
                print(f"      -> {s}")
            for j, nowa_nazwa, sciezka_kopii, _ in kopie:
                print(f"    [{j+1:>{len(str(n))}}/{n}] {nowa_nazwa}")
                try:
                    _przywroc_linki_w_xml(sciezka_kopii, LINKI_SERWEROWE)
                except Exception as exc:
                    _warn(f"Blad przywracania linkow w: {nowa_nazwa}\n"
                          f"        {type(exc).__name__}: {exc}", indent="    ")

    nazwy = [nazwa for (_, nazwa, _, _) in kopie]
    return nazwy, dane_kalibracji, dane_zakladek_per_kopia, klasa_wilg_per_kopia, zakresy_per_kopia


# =============================================================================
# GŁÓWNY PUNKT WEJŚCIA
# =============================================================================

def _main_impl():
    protokol = os.path.join(FOLDER, PROTOKOL_PLIK)
    szablon  = os.path.join(FOLDER, SZABLON_PLIK)

    SEP = "=" * 65

    # --- Walidacja plików wejściowych ---
    if not GENERUJ_EXCEL and not GENERUJ_WORD:
        print("[UWAGA] GENERUJ_EXCEL=False i GENERUJ_WORD=False — brak etapow do wykonania.")
        return

    if not os.path.exists(protokol):
        print(f"[BŁĄD] Plik protokołu nie istnieje:\n  {protokol}")
        return
    if GENERUJ_EXCEL and not os.path.exists(szablon):
        print(f"[BŁĄD] Plik szablonu nie istnieje:\n  {szablon}")
        return

    # --- Etapy 1-4: odczyt wszystkich danych przez xlwings (Excel COM) ---
    # xlwings uruchamia Excel i odczytuje prawdziwe wyliczone wartości formuł,
    # co rozwiązuje problem None przy data_only=True w openpyxl.
    print(SEP)
    print("ETAP 1-4  Odczyt danych z protokołu (przez Excel COM / xlwings)")
    print(SEP)
    print("  Uruchamiam Excel w tle — proszę czekać…")
    try:
        dane_s2, dane_zakladek, dane_ef, f24_per_kopia = wczytaj_wszystko_xlwings(
            protokol,
            ark_s2=ARKUSZ_STRONA2,
            ark_s3=ARKUSZ_STRONA3,
            start_s2=START_ROW_S2,
            start_s3=START_ROW_S3,
            blok=BLOK_S3,
            start_col_e=START_COL_E_S3,
            start_col_f=START_COL_F_S3,
            krok=KROK_COL_EF,
        )
    except Exception as e:
        print(f"[BŁĄD] Odczyt przez xlwings nie powiódł się: {e}")
        return

    print(f"  Znaleziono kopii: {len(dane_s2)}")
    if not dane_s2:
        print("[BŁĄD] Brak danych w Strona 2 — skrypt kończy pracę.")
        return

    # Przyrzady wyszarzone w calosci na Stronie 3 nie maja po co dostawac kopii
    # ani swiadectwa — wypadaja z obiegu tutaj, zanim ruszy Excel.
    if POMIJAJ_PRZYRZADY_BEZ_DANYCH:
        dane_s2, dane_ef, f24_per_kopia, pominiete = _odfiltruj_przyrzady_bez_danych(
            dane_s2, dane_ef, f24_per_kopia)
        if pominiete:
            print(f"  Pomijam {len(pominiete)} przyrzad(ow) bez aktywnych pomiarow "
                  f"(wszystkie bloki Strona 3 wyszarzone):")
            for nr, rekord in pominiete:
                print(f"    • przyrzad {nr}: {_numer_do_nazwy(rekord) or '(brak nr)'}")
            print(f"  Do przetworzenia zostaje: {len(dane_s2)} przyrzad(ow).")
        if not dane_s2:
            print("[BŁĄD] Zaden przyrzad nie ma aktywnych pomiarow na Stronie 3 — "
                  "nie ma czego generowac. Sprawdz kolory komorek E/F.")
            return

    if not dane_zakladek:
        print("[OSTRZEŻENIE] Brak definicji zakładek w Strona 3.")
    else:
        print(f"  Zakładki ({len(dane_zakladek)}):")
        for zd in dane_zakladek:
            print(f"    • {zd['nazwa']}")

    print(f"  Wczytano dane E/F: {len(dane_ef)} kopii × {len(dane_zakladek)} zakładek.")
    print(f"  Wczytano dane F24 (Wyniki): {len(f24_per_kopia)} wartości.")

    dane_zakladek_per_kopia = _zbuduj_dane_zakladek_per_kopia(dane_zakladek, dane_ef)
    nazwy = []
    dane_kalibracji = []
    klasa_wilg_per_kopia = []
    zakresy_per_kopia = []

    # --- Tworzenie i przetwarzanie kopii Excel ---
    if GENERUJ_EXCEL:
        print(SEP)
        print("ETAP 1+2+3+4  Tworzenie kopii i dostosowanie zakładek")
        print(SEP)

        nazwy, dane_kalibracji, dane_zakladek_per_kopia, klasa_wilg_per_kopia, zakresy_per_kopia = utworz_kopie(
            FOLDER, SZABLON_PLIK,
            dane_s2, dane_zakladek, dane_ef,
            ARKUSZ_WYNIKI,
            f24_per_kopia,
        )

        if USUWAJ_PUSTE_BLOKI_KOPII_S3:
            liczby_zakladek = [len(z) for z in dane_zakladek_per_kopia]
            if liczby_zakladek:
                min_z = min(liczby_zakladek)
                max_z = max(liczby_zakladek)
                print(f"  Aktywne zakladki robocze per kopia: min={min_z}, max={max_z}.")

        print(SEP)
        print(
            f"Zakończono etapy Excel. Utworzono {len(nazwy)} kopii. "
            f"Liczba zakładek roboczych jest dobierana per kopia + {ARKUSZ_WYNIKI}."
        )
        print(SEP)
    else:
        print(SEP)
        print("ETAP Excel pominięty (GENERUJ_EXCEL=False)")
        print(SEP)

        nazwy = [generuj_nazwe_pliku(SZABLON_PLIK, rekord["O"], _numer_do_nazwy(rekord),
                                     nr_przyrzadu=rekord.get("_nr_przyrzadu", j + 1))
                 for j, rekord in enumerate(dane_s2)]
        brakujace = [nazwa for nazwa in nazwy if not os.path.exists(os.path.join(FOLDER, nazwa))]
        if brakujace:
            print("[BŁĄD] Brak gotowych kopii Excel wymaganych do dalszych etapów:")
            for nazwa in brakujace:
                print(f"  - {nazwa}")
            return

        print("  Odczyt kalibracji z istniejacych kopii Excel...")
        dane_kalibracji, klasa_wilg_per_kopia = _odczytaj_kalibracje_dla_istniejacych_kopii(
            FOLDER,
            nazwy,
            ARKUSZ_WYNIKI,
        )

        if GENERUJ_WORD and (SZABLON_WORD_TYLKO_TEMP or SZABLON_WORD_Z_RH or SZABLON_WORD_MIESZANY):
            print("  Obliczam warunki srodowiskowe z istniejacych kopii...")
            zakresy_per_kopia = _oblicz_zakresy_srodowiskowe_z_istniejacych_kopii(
                FOLDER, nazwy, dane_zakladek_per_kopia)

    # --- Etap 7: dokumenty Word (świadectwa wzorcowania) ---
    if not GENERUJ_WORD:
        print("[INFO] Etap 7 pominięty (GENERUJ_WORD=False).")
    elif SZABLON_WORD_TYLKO_TEMP or SZABLON_WORD_Z_RH or SZABLON_WORD_MIESZANY:
        print(SEP)
        print("ETAP 7  Tworzenie świadectw wzorcowania (dokumenty Word)")
        print(SEP)
        nazwy_word = utworz_kopie_word(
            FOLDER, SZABLON_WORD_TYLKO_TEMP, SZABLON_WORD_Z_RH, SZABLON_WORD_MIESZANY,
            dane_s2, nazwy,
            dane_zakladek, dane_kalibracji, NR_SW_POCZATKOWY,
            dane_zakladek_per_kopia=dane_zakladek_per_kopia,
            klasa_wilg_per_kopia=klasa_wilg_per_kopia,
            warunki_per_kopia=zakresy_per_kopia,
        )
        print(SEP)
        print(f"Zakończono Etap 7. Utworzono {len(nazwy_word)} dokumentów Word.")
        print(SEP)

def _excel_options_subkey():
    """Zwraca podklucz rejestru Excel\\Options dla zainstalowanej wersji Office."""
    import winreg
    for wer in ("16.0", "15.0", "14.0"):
        sub = rf"Software\Microsoft\Office\{wer}\Excel\Options"
        try:
            with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub):
                return sub
        except OSError:
            continue
    return None


def _czytaj_autorecover():
    """Globalne ustawienie AutoRecover Excela: True/False (None = nie udalo sie odczytac)."""
    import winreg
    sub = _excel_options_subkey()
    if not sub:
        return None
    try:
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub) as k:
            val, _ = winreg.QueryValueEx(k, "AutoRecoverEnabled")
            return bool(val)
    except FileNotFoundError:
        return True   # brak wartosci = Excel domyslnie ma AutoRecover wlaczone
    except OSError:
        return None


def _ustaw_autorecover(wlaczone):
    """Ustawia globalne AutoRecover Excela (rejestr). Cicho ignoruje bledy."""
    import winreg
    sub = _excel_options_subkey()
    if not sub:
        return
    try:
        with winreg.OpenKey(winreg.HKEY_CURRENT_USER, sub, 0, winreg.KEY_SET_VALUE) as k:
            winreg.SetValueEx(k, "AutoRecoverEnabled", 0, winreg.REG_DWORD, 1 if wlaczone else 0)
    except OSError:
        pass


def _rozmiar_sciezki(p):
    """Rozmiar pliku lub calego folderu w bajtach (bledy -> 0)."""
    if os.path.isfile(p):
        try:
            return os.path.getsize(p)
        except OSError:
            return 0
    total = 0
    for root, _dirs, files in os.walk(p):
        for f in files:
            try:
                total += os.path.getsize(os.path.join(root, f))
            except OSError:
                pass
    return total


def _wyczysc_autorecover_folder():
    """
    Usuwa stare pliki autoodzyskiwania Excela z %AppData%\\Microsoft\\Excel,
    zostawiajac XLSTART (dodatki) oraz pliki *.xlb (ustawienia paskow narzedzi).
    Usuwa tylko elementy starsze niz CZYSC_AUTORECOVER_STARSZE_NIZ_DNI
    (0 = bez ograniczenia wieku). Pliki zablokowane (otwarty Excel) sa pomijane.
    """
    folder = os.path.join(os.environ.get("APPDATA", ""), "Microsoft", "Excel")
    if not os.path.isdir(folder):
        return

    granica = None
    if CZYSC_AUTORECOVER_STARSZE_NIZ_DNI and CZYSC_AUTORECOVER_STARSZE_NIZ_DNI > 0:
        granica = datetime.datetime.now() - datetime.timedelta(days=CZYSC_AUTORECOVER_STARSZE_NIZ_DNI)

    usuniete = 0
    zwolnione = 0
    pominiete_swieze = 0
    for nazwa in os.listdir(folder):
        if nazwa.upper() == "XLSTART" or nazwa.lower().endswith(".xlb"):
            continue
        pelna = os.path.join(folder, nazwa)
        try:
            mtime = datetime.datetime.fromtimestamp(os.path.getmtime(pelna))
        except OSError:
            continue
        if granica is not None and mtime >= granica:
            pominiete_swieze += 1
            continue  # zbyt swieze — moze byc potrzebne (reczne odzyskiwanie)
        try:
            rozmiar = _rozmiar_sciezki(pelna)
            if os.path.isdir(pelna):
                shutil.rmtree(pelna, ignore_errors=True)
            else:
                os.remove(pelna)
            usuniete += 1
            zwolnione += rozmiar
        except OSError:
            pass  # plik zablokowany przez otwarty Excel — pomijamy

    if usuniete:
        print(f"  [AutoRecover] Usunieto {usuniete} starych elementow odzyskiwania "
              f"({zwolnione // (1024 * 1024)} MB).")
    if pominiete_swieze:
        print(f"  [AutoRecover] Zachowano {pominiete_swieze} swiezych elementow "
              f"(< {CZYSC_AUTORECOVER_STARSZE_NIZ_DNI} dni — mozliwe reczne odzyskiwanie).")


def main():
    # Skrypt wylacza AutoRecover tylko NA CZAS swojej pracy (instancje Excela nie
    # tworza plikow .xar — nawet przy ewentualnej awarii nie zostaje smieci w
    # %AppData%\Microsoft\Excel). Preferencje uzytkownika dla RECZNEJ pracy w
    # Excelu zapamietujemy tutaj i przywracamy w finally.
    _ar_user = _czytaj_autorecover()

    # Sprzatanie nagromadzonych plikow autoodzyskiwania PRZED uruchomieniem Excela
    # (inaczej Excel probuje je odzyskac przy starcie i potrafi paść).
    if CZYSC_AUTORECOVER:
        try:
            _wyczysc_autorecover_folder()
        except Exception as e:
            print(f"  [AutoRecover] Sprzatanie pominieto: {type(e).__name__}: {e}")

    try:
        _main_impl()
    finally:
        if _ar_user is not None:
            _ustaw_autorecover(_ar_user)


if __name__ == "__main__":
    main()
